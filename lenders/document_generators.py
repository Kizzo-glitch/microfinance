
from datetime import timezone
from docx import Document

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from io import BytesIO



class CBLDocumentGenerator:
    """Generate CBL-required documents"""
    
    def __init__(self, registration):
        self.registration = registration
    
    def generate_aml_manual(self):
        """Generate AML/CFT Manual compliant with Money Laundering Act 2008"""
        doc = Document()
        
        doc.add_heading(f'{self.registration.company_name}', 0)
        doc.add_heading('ANTI-MONEY LAUNDERING AND COUNTER-FINANCING OF TERRORISM', 1)
        doc.add_heading('(AML/CFT) POLICY MANUAL', 1)
        
        doc.add_paragraph(f'Version 1.0 | {timezone.now().strftime("%B %Y")}')
        doc.add_page_break()
        
        # Table of Contents
        doc.add_heading('TABLE OF CONTENTS', 1)
        toc_items = [
            '1. Introduction',
            '2. Legal and Regulatory Framework',
            '3. AML/CFT Risk Assessment',
            '4. Customer Due Diligence (CDD)',
            '5. Enhanced Due Diligence (EDD)',
            '6. Politically Exposed Persons (PEPs)',
            '7. Transaction Monitoring',
            '8. Suspicious Activity Reporting',
            '9. Record Keeping',
            '10. Staff Training',
            '11. Internal Controls and Audit',
            '12. Reporting to FIU',
        ]
        for item in toc_items:
            doc.add_paragraph(item, style='List Number')
        
        doc.add_page_break()
        
        # Chapter 1: Introduction
        doc.add_heading('1. INTRODUCTION', 1)
        doc.add_paragraph(f"""
        This Anti-Money Laundering and Counter-Financing of Terrorism (AML/CFT) Policy Manual 
        establishes {self.registration.company_name}'s commitment to preventing the use of our 
        services for money laundering or terrorist financing activities.
        
        This manual is prepared in compliance with:
        - Money Laundering and Proceeds of Crime Act 2008 as amended 2016
        - Money Laundering and Proceeds of Crime Regulations 2017
        - Financial Institutions Act 2012
        - Central Bank of Lesotho AML/CFT Guidelines
        """)
        
        # Chapter 2: Legal Framework
        doc.add_heading('2. LEGAL AND REGULATORY FRAMEWORK', 1)
        doc.add_paragraph("""
        2.1 Money Laundering and Proceeds of Crime Act 2008
        
        This Act criminalizes money laundering and establishes obligations for financial 
        institutions to implement AML/CFT measures.
        
        2.2 Key Obligations
        - Customer Due Diligence (CDD)
        - Record Keeping (minimum 7 years)
        - Suspicious Activity Reporting to FIU
        - Staff Training
        - Internal Controls
        """)
        
        # Chapter 3: Risk Assessment
        doc.add_heading('3. AML/CFT RISK ASSESSMENT', 1)
        doc.add_paragraph("""
        3.1 Risk-Based Approach
        
        {self.registration.company_name} adopts a risk-based approach to AML/CFT, focusing 
        resources on higher-risk customers and transactions.
        
        3.2 Risk Factors
        
        Customer Risk Factors:
        - Occupation (cash-intensive businesses)
        - Geographic location (high-risk areas)
        - Transaction patterns (unusual or complex)
        - Politically Exposed Persons (PEPs)
        
        Product/Service Risk:
        - Large loans (>M 50,000)
        - Cross-border transactions
        - Cash-intensive products
        
        3.3 Risk Categorization
        
        LOW RISK:
        - Employed individuals with verified income
        - Small loans (<M 10,000)
        - Group lending with social collateral
        
        MEDIUM RISK:
        - Self-employed individuals
        - Loans M 10,000 - M 50,000
        - New customers
        
        HIGH RISK:
        - Cash-intensive businesses
        - Politically Exposed Persons
        - Loans >M 50,000
        - Non-resident customers
        """)
        
        # Chapter 4: Customer Due Diligence
        doc.add_heading('4. CUSTOMER DUE DILIGENCE (CDD)', 1)
        doc.add_paragraph("""
        4.1 When to Conduct CDD
        
        CDD must be conducted:
        - When establishing a business relationship
        - When carrying out occasional transactions above M 25,000
        - When there is suspicion of money laundering or terrorist financing
        - When there are doubts about previously obtained identification data
        
        4.2 CDD Measures
        
        Standard CDD includes:
        a) Identification and verification of customer identity
           - National ID or passport
           - Proof of address (utility bill, bank statement, chief's letter)
        
        b) Verification of source of funds
           - Payslips for employed
           - Business records for self-employed
           - Bank statements
        
        c) Understanding purpose and nature of business relationship
           - Loan purpose
           - Expected transaction activity
        
        d) Ongoing monitoring of business relationship
           - Review of transactions
           - Update of customer information
        
        4.3 Documentation Required
        
        For Individual Customers:
        ✓ National ID or passport (certified copy)
        ✓ Proof of address (within 3 months)
        ✓ Proof of income (payslips, business records)
        ✓ Tax clearance (if self-employed)
        ✓ Loan application form
        
        For Business Customers:
        ✓ Company registration certificate
        ✓ Memorandum and articles of association
        ✓ Tax clearance certificate
        ✓ Directors' identification documents
        ✓ Beneficial ownership declaration
        ✓ Business financial statements
        """)
        
        # [Continue with remaining chapters...]
        
        return doc
    
    def generate_risk_management_manual(self):
        """Generate Risk Management Manual for lending operations"""
        # Similar comprehensive template
        pass
    
    def generate_consumer_complaints_procedure(self):
        """Generate Consumer Complaints and Redress Procedure"""
        doc = Document()
        
        doc.add_heading(f'{self.registration.company_name}', 0)
        doc.add_heading('CONSUMER COMPLAINTS AND REDRESS PROCEDURE', 1)
        doc.add_paragraph(f'Effective Date: {timezone.now().strftime("%d %B %Y")}')
        
        doc.add_page_break()
        
        doc.add_heading('1. PURPOSE', 1)
        doc.add_paragraph("""
        This procedure establishes a fair, accessible, and efficient process for handling 
        customer complaints in compliance with Regulation 11 of the Microfinance Institutions 
        Regulations 2014.
        """)
        
        doc.add_heading('2. SCOPE', 1)
        doc.add_paragraph("""
        This procedure applies to all complaints received from customers regarding our 
        products, services, or staff conduct.
        """)
        
        doc.add_heading('3. COMPLAINTS CHANNELS', 1)
        doc.add_paragraph("""
        Customers can submit complaints through:
        - In-person at any branch
        - Phone: [PHONE NUMBER]
        - Email: complaints@[company].com
        - WhatsApp: [NUMBER]
        - Written letter to registered address
        - Online complaint form
        """)
        
        doc.add_heading('4. COMPLAINT HANDLING PROCESS', 1)
        doc.add_paragraph("""
        STAGE 1: ACKNOWLEDGMENT (Within 24 hours)
        - Complaint logged in complaints register
        - Unique reference number issued
        - Acknowledgment sent to customer
        
        STAGE 2: INVESTIGATION (Within 5 business days)
        - Complaint assigned to responsible officer
        - Facts gathered from all parties
        - Review of relevant documents and records
        - Determine root cause
        
        STAGE 3: RESOLUTION (Within 14 days)
        - Proposed resolution developed
        - Resolution communicated to customer
        - Customer acceptance obtained
        - Remedial action implemented
        
        STAGE 4: ESCALATION (If unresolved)
        - Customer may escalate to senior management
        - Review by complaints committee
        - Final decision within 7 days
        
        STAGE 5: EXTERNAL RESOLUTION
        - Customer may refer to Central Bank of Lesotho
        - Contact: Consumer Protection Department, CBL
        - Phone: +266 2231 4281
        - Email: info@centralbank.org.ls
        """)
        
        return doc
    

### Phase 3: Document Generation Engine
class CBLApplicationGenerator:
    """Auto-generate CBL application documents"""
    
    def __init__(self, registration):
        self.registration = registration
    
    def generate_complete_package(self):
        """Generate all required documents"""
        
        documents = []
        
        # 1. Cover Letter to CBL
        documents.append(self.generate_cover_letter())
        
        # 2. Application Form (Form MFI-001)
        documents.append(self.generate_application_form())
        
        # 3. Business Plan
        documents.append(self.generate_business_plan())
        
        # 4. Financial Projections
        documents.append(self.generate_financial_projections())
        
        # 5. AML/CFT Policy Manual
        documents.append(self.generate_aml_manual())
        
        # 6. Operations Manual
        documents.append(self.generate_operations_manual())
        
        # 7. Organizational Structure
        documents.append(self.generate_org_structure())
        
        # 8. Board Resolution
        documents.append(self.generate_board_resolution())
        
        return documents
    
    def generate_cover_letter(self):
        """Generate formal cover letter to CBL"""
        
        doc = Document()
        
        # Header
        doc.add_heading(f'{self.registration.company_name}', 0)
        doc.add_paragraph(f'{self.registration.applicant.email}')
        doc.add_paragraph(f'{timezone.now().strftime("%d %B %Y")}')
        
        doc.add_paragraph('')
        doc.add_paragraph('The Governor')
        doc.add_paragraph('Central Bank of Lesotho')
        doc.add_paragraph('P.O. Box 1184')
        doc.add_paragraph('Maseru 100, Lesotho')
        
        doc.add_paragraph('')
        doc.add_heading('RE: APPLICATION FOR MICROFINANCE INSTITUTION LICENSE', 1)
        
        content = f"""
        Dear Sir/Madam,
        
        We hereby submit our application for a Microfinance Institution License under the 
        Financial Institutions Act. {self.registration.company_name} intends to provide 
        microfinance services to underserved communities in Lesotho, with initial capital 
        of M {self.registration.available_capital:,.2f}.
        
        This application package includes:
        1. Completed application form (Form MFI-001)
        2. Comprehensive business plan
        3. Three-year financial projections
        4. AML/CFT policy manual
        5. Operations manual
        6. Organizational structure and governance framework
        7. Board member details and CVs
        8. Proof of minimum capital
        9. All supporting documents as required
        
        We have demonstrated our commitment to regulatory compliance by utilizing the 
        [Your Platform Name] compliance platform, which ensures ongoing adherence to 
        CBL regulations.
        
        We look forward to your favorable consideration and remain available for any 
        clarifications or additional information you may require.
        
        Yours faithfully,
        
        
        _______________________________
        {self.registration.applicant.get_full_name()}
        {self.registration.company_name}
        """
        
        doc.add_paragraph(content)
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return ('cover_letter.docx', buffer)
    
    def generate_business_plan(self):
        """Generate comprehensive business plan"""
        
        doc = Document()
        
        # Title Page
        doc.add_heading(f'BUSINESS PLAN', 0)
        doc.add_heading(f'{self.registration.company_name}', 1)
        doc.add_heading(f'Microfinance Institution', 2)
        doc.add_paragraph(f'{timezone.now().year} - {timezone.now().year + 3}')
        
        # Executive Summary
        doc.add_page_break()
        doc.add_heading('EXECUTIVE SUMMARY', 1)
        doc.add_paragraph(f"""
        {self.registration.company_name} is a [new/existing] microfinance institution 
        established to provide accessible, affordable credit to unbanked and underbanked 
        individuals in Lesotho. With initial capital of M {self.registration.available_capital:,.2f}, 
        we aim to serve [target number] borrowers in our first year of operations.
        
        Our unique value proposition includes:
        - Group-based lending model reducing risk
        - Digital-first platform for efficiency
        - Focus on women entrepreneurs and rural communities
        - Comprehensive financial literacy programs
        - Partnership with [Your Platform Name] for technology and compliance
        """)
        
        # Market Analysis
        doc.add_heading('MARKET ANALYSIS', 1)
        doc.add_paragraph("""
        Lesotho Financial Inclusion Landscape:
        - Total adult population: 1.2 million
        - Unbanked: 60% (720,000 individuals)
        - Informal credit users: 450,000
        - Microfinance penetration: <15%
        
        Target Market:
        - Primary: Rural employed (250,000)
        - Secondary: Urban informal workers (200,000)
        - Focus: Women entrepreneurs (55% of portfolio)
        """)
        
        # Products and Services
        doc.add_heading('PRODUCTS AND SERVICES', 1)
        # ... continue with all business plan sections
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return ('business_plan.docx', buffer)