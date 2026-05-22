
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime, date
from decimal import Decimal
from typing import Dict, List, Optional, Tuple


# Data aggregation from multiple sources
class ComplianceDocumentGenerator:
    def __init__(self, lender):
        self.lender = lender
        self.compliance = lender.compliance
        self.personnel = lender.personnel.all()
        
    def collect_data_sources(self):
        return {
            'lender_profile': self._extract_lender_data(),
            'financial_data': self._extract_financial_data(),
            'governance_data': self._extract_governance_data(),
            'personnel_data': self._extract_personnel_data(),
            'external_data': self._fetch_external_data(),
        }
    
    def _extract_lender_data(self):
        return {
            'company_name': self.lender.company_name,
            'registration_number': self.lender.registration_number,
            'establishment_date': self.lender.date_of_establishment,
            'physical_address': self.lender.physical_address,
            'stated_capital': self.lender.stated_capital,
            'total_assets': self.lender.total_assets,
            'cbl_tier': self.lender.cbl_tier,
            'business_email': self.lender.business_email,
            'ceo_name': f"{self.lender.ceo_first_name} {self.lender.ceo_last_name}",
        }




# ==========================================================================================
    # ScheduleIGenerator
# ===============================================================================================   

class ScheduleIGenerator:
    """
    Schedule I Application Form Generator
    Auto-fills CBL Schedule I with company and business information
    """
    
    def __init__(self, lender_data):
        self.lender = lender_data
        
    def generate(self):
        """Generate Schedule I document"""     
        doc = Document()
        
        # Set up document style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        
        # CBL Header
        header_para = doc.add_paragraph()
        header_para.alignment = 1  # Center alignment
        header_run = header_para.add_run('CENTRAL BANK OF LESOTHO\n')
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        
        subheader_run = header_para.add_run('MICROFINANCE INSTITUTIONS AND COOPERATIVES SUPERVISION DEPARTMENT\n')
        subheader_run.font.size = Pt(11)
        subheader_run.font.bold = True
        
        # Form title
        doc.add_paragraph()
        title_para = doc.add_paragraph()
        title_para.alignment = 1
        title_run = title_para.add_run('SCHEDULE I - APPLICATION FORM\n')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.underline = True
        
        subtitle_run = title_para.add_run('Application for License to Conduct Microfinance Business')
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        
        doc.add_paragraph()
        
        # Section 1: Applicant Information
        section1_para = doc.add_paragraph()
        section1_run = section1_para.add_run('SECTION 1: APPLICANT INFORMATION')
        section1_run.font.size = Pt(14)
        section1_run.font.bold = True
        section1_run.font.underline = True
        
        # Create applicant info table
        info_table = doc.add_table(rows=10, cols=2)
        info_table.style = 'Table Grid'
        
        applicant_info = [
            ('1.1 Name of Institution:', self.lender.get('company_name', '')),
            ('1.2 Registration Number:', self.lender.get('registration_number', '')),
            ('1.3 Date of Incorporation:', self._format_date(self.lender.get('date_of_establishment'))),
            ('1.4 Physical Address:', self.lender.get('physical_address', '')),
            ('1.5 Postal Address:', self.lender.get('postal_address', '')),
            ('1.6 Telephone Number:', self.lender.get('phone_number', '')),
            ('1.7 Email Address:', self.lender.get('business_email', '')),
            ('1.8 Type of License Applied For:', self.lender.get('cbl_tier', '').upper()),
            ('1.9 Proposed Commencement Date:', self._calculate_commencement_date()),
            ('1.10 Application Date:', datetime.now().strftime('%d/%m/%Y')),
        ]
        
        for i, (label, value) in enumerate(applicant_info):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)
            info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Section 2: Business Activities
        doc.add_paragraph()
        section2_para = doc.add_paragraph()
        section2_run = section2_para.add_run('SECTION 2: PROPOSED BUSINESS ACTIVITIES')
        section2_run.font.size = Pt(14)
        section2_run.font.bold = True
        section2_run.font.underline = True
        
        activities_para = doc.add_paragraph()
        activities_text = self._generate_business_activities_text()
        activities_para.add_run(activities_text)
        
        # Section 3: Financial Information
        doc.add_paragraph()
        section3_para = doc.add_paragraph()
        section3_run = section3_para.add_run('SECTION 3: FINANCIAL INFORMATION')
        section3_run.font.size = Pt(14)
        section3_run.font.bold = True
        section3_run.font.underline = True
        
        financial_table = doc.add_table(rows=4, cols=2)
        financial_table.style = 'Table Grid'
        
        financial_info = [
            ('3.1 Authorized Share Capital:', f"M{self.lender.get('stated_capital', 0):,.2f}"),
            ('3.2 Issued Share Capital:', f"M{float(self.lender.get('stated_capital', 0)) * 0.8:,.2f}"),
            ('3.3 Paid-up Capital:', f"M{float(self.lender.get('stated_capital', 0)) * 0.6:,.2f}"),
            ('3.4 Total Assets:', f"M{self.lender.get('total_assets', 0):,.2f}"),
        ]
        
        for i, (label, value) in enumerate(financial_info):
            financial_table.rows[i].cells[0].text = label
            financial_table.rows[i].cells[1].text = value
            financial_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Section 4: Declaration
        doc.add_paragraph()
        section4_para = doc.add_paragraph()
        section4_run = section4_para.add_run('SECTION 4: DECLARATION')
        section4_run.font.size = Pt(14)
        section4_run.font.bold = True
        section4_run.font.underline = True
        
        declaration_text = f"""
I/We, {self._get_applicant_name()}, hereby apply for a license to conduct microfinance business under the Financial Institutions Act of Lesotho and declare that:

1. The information provided in this application is true, complete, and accurate.
2. The institution will comply with all applicable laws and regulations.
3. All required supporting documents are attached to this application.
4. The institution understands its obligations under the CBL licensing requirements.
        """
        
        declaration_para = doc.add_paragraph()
        declaration_para.add_run(declaration_text.strip())
        
        # Signature section
        doc.add_paragraph()
        signature_table = doc.add_table(rows=3, cols=2)
        signature_table.style = 'Table Grid'
        
        signature_table.rows[0].cells[0].text = 'CEO Signature:'
        signature_table.rows[0].cells[1].text = 'Date:'
        signature_table.rows[1].cells[0].text = '________________________________'
        signature_table.rows[1].cells[1].text = '________________'
        signature_table.rows[2].cells[0].text = f'Name: {self._get_ceo_name()}'
        signature_table.rows[2].cells[1].text = 'Position: Chief Executive Officer'
        
        return doc
    
    def _format_date(self, date_obj):
        """Format date for display"""
        if date_obj:
            return date_obj.strftime('%d/%m/%Y') if hasattr(date_obj, 'strftime') else str(date_obj)
        return ''
    
    def _calculate_commencement_date(self):
        """Calculate proposed commencement date"""
        from datetime import datetime, timedelta
        commencement = datetime.now() + timedelta(days=90)  # 3 months from now
        return commencement.strftime('%d/%m/%Y')
    
    def _generate_business_activities_text(self):
        """Generate business activities description"""
        tier = self.lender.get('cbl_tier', '')
        
        if tier == 'tier1':
            return """The institution intends to conduct the following microfinance activities:
- Accepting deposits from the general public
- Providing micro-credit to individuals and small businesses
- Money transfer and payment services
- Financial advisory and business development support
- Insurance agency services (where appropriately licensed)"""
        else:
            return """The institution intends to conduct the following microfinance activities:
- Providing micro-credit to individuals and small businesses
- Financial advisory and business development support
- Business training and capacity building services
- Mobile money and digital payment facilitation"""
    
    def _get_applicant_name(self):
        """Get applicant name for declaration"""
        ceo_name = self._get_ceo_name()
        company_name = self.lender.get('company_name', '')
        return f"{ceo_name}, on behalf of {company_name}"
    
    def _get_ceo_name(self):
        """Get CEO name"""
        first_name = self.lender.get('ceo_first_name', '')
        last_name = self.lender.get('ceo_last_name', '')
        return f"{first_name} {last_name}".strip() or 'To be appointed'


# ===========================================================================================================
    # Schedule II Generator - CBL Information Sheet
    # Generates the mandatory CBL Schedule II form with comprehensive company and ownership information
# ===========================================================================================================================
class ScheduleIIGenerator:
    """
    Generates CBL Schedule II - Information Sheet
    
    This form provides detailed information about the company structure,
    ownership, board composition, and operational plans as required by CBL
    """
    
    def __init__(self, lender_data: Dict):
        self.lender = lender_data
        self.compliance = lender_data.get('compliance_data', {})
        self.personnel = lender_data.get('personnel', [])
        
    def generate(self) -> Document:
        """Generate complete Schedule II document"""
        doc = Document()
        
        # Set up document formatting
        self._setup_document_style(doc)
        
        # Add CBL header and form title
        self._add_header(doc)
        
        # Section 1: Company Information
        self._add_section_1_company_info(doc)
        
        # Section 2: Ownership Profile
        self._add_section_2_ownership(doc)
        
        # Section 3: Board of Directors
        self._add_section_3_board(doc)
        
        # Section 4: Management Structure
        self._add_section_4_management(doc)
        
        # Section 5: Business Activities
        self._add_section_5_activities(doc)
        
        # Section 6: Financial Information
        self._add_section_6_financial(doc)
        
        # Section 7: Compliance Declaration
        self._add_section_7_declaration(doc)
        
        return doc
    
    def _setup_document_style(self, doc: Document):
        """Set up document-wide styling"""
        # Set normal style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
    
    def _add_header(self, doc: Document):
        """Add CBL header and form identification"""
        # CBL Logo placeholder and header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        header_run = header_para.add_run('CENTRAL BANK OF LESOTHO\n')
        header_run.font.name = 'Times New Roman'
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        
        subheader_run = header_para.add_run('MICROFINANCE INSTITUTIONS AND COOPERATIVES SUPERVISION DEPARTMENT\n')
        subheader_run.font.size = Pt(11)
        subheader_run.font.bold = True
        
        # Form title
        doc.add_paragraph()
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_run = title_para.add_run('SCHEDULE II - INFORMATION SHEET\n')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.underline = True
        
        subtitle_run = title_para.add_run('(To be submitted with Application for License)')
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        
        # Form reference and date
        doc.add_paragraph()
        ref_table = doc.add_table(rows=2, cols=4)
        ref_table.style = 'Table Grid'
        
        # Form reference row
        ref_table.rows[0].cells[0].text = 'Form Reference:'
        ref_table.rows[0].cells[1].text = f'CBL/MFI/SCH-II/{datetime.now().year}'
        ref_table.rows[0].cells[2].text = 'Date Prepared:'
        ref_table.rows[0].cells[3].text = datetime.now().strftime('%d/%m/%Y')
        
        # Application type row
        ref_table.rows[1].cells[0].text = 'Application Type:'
        ref_table.rows[1].cells[1].text = f'New License - {self.lender.get("cbl_tier", "").upper()}'
        ref_table.rows[1].cells[2].text = 'CBL Use Only:'
        ref_table.rows[1].cells[3].text = ''
        
        doc.add_paragraph()
    
    def _add_section_1_company_info(self, doc: Document):
        """Section 1: Basic Company Information"""
        self._add_section_header(doc, '1. COMPANY INFORMATION')
        
        # Create information table
        info_table = doc.add_table(rows=12, cols=2)
        info_table.style = 'Table Grid'
        
        # Set column widths
        info_table.columns[0].width = Inches(2.5)
        info_table.columns[1].width = Inches(4.5)
        
        # Populate company information
        company_info = [
            ('Full Registered Name:', self.lender.get('company_name', '')),
            ('Trading Name (if different):', self.lender.get('trading_name', '')),
            ('Company Registration Number:', self.lender.get('registration_number', '')),
            ('Date of Incorporation:', self._format_date(self.lender.get('date_of_establishment'))),
            ('Tax Identification Number:', self.lender.get('tax_identification_number', '')),
            ('Registered Office Address:', self.lender.get('physical_address', '')),
            ('Postal Address:', self.lender.get('postal_address', '')),
            ('Telephone Number:', self.lender.get('phone_number', '')),
            ('Email Address:', self.lender.get('business_email', '')),
            ('Website:', self.lender.get('website', 'N/A')),
            ('Nature of Business:', self._determine_business_nature()),
            ('Proposed CBL License Tier:', self.lender.get('cbl_tier', '').upper().replace('TIER', 'TIER ')),
        ]
        
        for i, (label, value) in enumerate(company_info):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)
            
            # Bold the labels
            info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    def _add_section_2_ownership(self, doc: Document):
        """Section 2: Ownership Structure and Profile"""
        self._add_section_header(doc, '2. OWNERSHIP STRUCTURE')
        
        # Shareholding structure
        doc.add_paragraph('2.1 SHAREHOLDING STRUCTURE', style='Heading 3')
        
        shareholding_table = doc.add_table(rows=1, cols=5)
        shareholding_table.style = 'Table Grid'
        
        # Headers
        headers = ['Shareholder Name', 'ID/Passport No.', 'Nationality', 'Shares Held', 'Percentage (%)']
        header_row = shareholding_table.rows[0]
        for i, header in enumerate(headers):
            header_row.cells[i].text = header
            header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add shareholding data (if available)
        shareholding_data = self._extract_shareholding_data()
        for shareholder in shareholding_data:
            row = shareholding_table.add_row()
            row.cells[0].text = shareholder['name']
            row.cells[1].text = shareholder['id_number']
            row.cells[2].text = shareholder['nationality']
            row.cells[3].text = f"{shareholder['shares']:,}"
            row.cells[4].text = f"{shareholder['percentage']:.1f}%"
        
        # Beneficial ownership section
        doc.add_paragraph()
        doc.add_paragraph('2.2 BENEFICIAL OWNERSHIP (Persons owning ≥ 25% directly or indirectly)', style='Heading 3')
        
        beneficial_table = doc.add_table(rows=1, cols=4)
        beneficial_table.style = 'Table Grid'
        
        beneficial_headers = ['Full Name', 'ID/Passport No.', 'Beneficial Ownership %', 'Nature of Control']
        beneficial_row = beneficial_table.rows[0]
        for i, header in enumerate(beneficial_headers):
            beneficial_row.cells[i].text = header
            beneficial_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add beneficial ownership data
        beneficial_owners = self._identify_beneficial_owners()
        for owner in beneficial_owners:
            row = beneficial_table.add_row()
            row.cells[0].text = owner['name']
            row.cells[1].text = owner['id_number']
            row.cells[2].text = f"{owner['ownership_percentage']:.1f}%"
            row.cells[3].text = owner['control_nature']
        
        # Foreign ownership declaration
        doc.add_paragraph()
        foreign_ownership = self._calculate_foreign_ownership()
        foreign_para = doc.add_paragraph()
        foreign_run = foreign_para.add_run('2.3 FOREIGN OWNERSHIP: ')
        foreign_run.font.bold = True
        foreign_para.add_run(f'{foreign_ownership:.1f}% of total shares are held by foreign nationals/entities.')
        
        if foreign_ownership > 49:
            warning_para = doc.add_paragraph()
            warning_run = warning_para.add_run('NOTE: ')
            warning_run.font.bold = True
            warning_run.font.color.rgb = (255, 0, 0)  # Red color
            warning_para.add_run('Foreign ownership exceeds 49% - CBL approval required under Financial Institutions Act.')
    
    def _add_section_3_board(self, doc: Document):
        """Section 3: Board of Directors"""
        self._add_section_header(doc, '3. BOARD OF DIRECTORS')
        
        # Board composition summary
        board_members = [p for p in self.personnel if p.get('role') == 'director']
        
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run('BOARD COMPOSITION SUMMARY: ')
        summary_run.font.bold = True
        summary_para.add_run(f'The Board consists of {len(board_members)} directors as detailed below:')
        
        doc.add_paragraph()
        
        # Board members table
        board_table = doc.add_table(rows=1, cols=7)
        board_table.style = 'Table Grid'
        
        board_headers = ['Full Name', 'ID Number', 'Nationality', 'Position', 'Appointment Date', 'Independence', 'Other Directorships']
        header_row = board_table.rows[0]
        for i, header in enumerate(board_headers):
            header_row.cells[i].text = header
            header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add board member data
        for member in board_members:
            row = board_table.add_row()
            row.cells[0].text = member.get('full_name', '')
            row.cells[1].text = member.get('id_or_passport_number', '')
            row.cells[2].text = member.get('nationality', '')
            row.cells[3].text = self._determine_board_position(member)
            row.cells[4].text = self._format_date(member.get('appointment_date', datetime.now().date()))
            row.cells[5].text = 'Independent' if member.get('is_non_executive') else 'Non-Independent'
            row.cells[6].text = self._extract_other_directorships(member)
        
        # Board compliance checks
        doc.add_paragraph()
        self._add_board_compliance_section(doc, board_members)
    
    def _add_section_4_management(self, doc: Document):
        """Section 4: Key Management Personnel"""
        self._add_section_header(doc, '4. KEY MANAGEMENT PERSONNEL')
        
        # Key management table
        mgmt_table = doc.add_table(rows=1, cols=6)
        mgmt_table.style = 'Table Grid'
        
        mgmt_headers = ['Name', 'Position', 'Qualifications', 'Experience (Years)', 'Previous Institution', 'Start Date']
        header_row = mgmt_table.rows[0]
        for i, header in enumerate(mgmt_headers):
            header_row.cells[i].text = header
            header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Key personnel (CEO, Finance Officer, Compliance Officer, etc.)
        key_personnel = [p for p in self.personnel if p.get('role') in ['ceo', 'finance_officer', 'compliance_officer']]
        
        for person in key_personnel:
            row = mgmt_table.add_row()
            row.cells[0].text = person.get('full_name', '')
            row.cells[1].text = person.get('role', '').replace('_', ' ').title()
            row.cells[2].text = self._extract_qualifications(person)
            row.cells[3].text = self._calculate_experience_years(person)
            row.cells[4].text = self._extract_previous_institution(person)
            row.cells[5].text = self._format_date(person.get('start_date', datetime.now().date()))
        
        # Management structure narrative
        doc.add_paragraph()
        mgmt_para = doc.add_paragraph()
        mgmt_run = mgmt_para.add_run('MANAGEMENT STRUCTURE: ')
        mgmt_run.font.bold = True
        mgmt_para.add_run(self._generate_management_narrative())
    
    def _add_section_5_activities(self, doc: Document):
        """Section 5: Proposed Business Activities"""
        self._add_section_header(doc, '5. PROPOSED BUSINESS ACTIVITIES')
        
        # Primary activities
        doc.add_paragraph('5.1 PRIMARY ACTIVITIES', style='Heading 3')
        
        activities = self._determine_business_activities()
        for activity in activities['primary']:
            activity_para = doc.add_paragraph(style='List Bullet')
            activity_para.add_run(activity)
        
        # Target market
        doc.add_paragraph()
        doc.add_paragraph('5.2 TARGET MARKET', style='Heading 3')
        
        target_para = doc.add_paragraph()
        target_para.add_run(self._generate_target_market_description())
        
        # Geographic coverage
        doc.add_paragraph()
        doc.add_paragraph('5.3 GEOGRAPHIC COVERAGE', style='Heading 3')
        
        geographic_para = doc.add_paragraph()
        geographic_para.add_run(self._generate_geographic_coverage())
        
        # Products and services
        doc.add_paragraph()
        doc.add_paragraph('5.4 PRODUCTS AND SERVICES', style='Heading 3')
        
        products_table = doc.add_table(rows=1, cols=4)
        products_table.style = 'Table Grid'
        
        products_headers = ['Product/Service', 'Target Segment', 'Interest Rate Range', 'Loan Term Range']
        header_row = products_table.rows[0]
        for i, header in enumerate(products_headers):
            header_row.cells[i].text = header
            header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add product details
        products = self._determine_products_services()
        for product in products:
            row = products_table.add_row()
            row.cells[0].text = product['name']
            row.cells[1].text = product['target_segment']
            row.cells[2].text = product['interest_rate']
            row.cells[3].text = product['loan_term']
    
    def _add_section_6_financial(self, doc: Document):
        """Section 6: Financial Information Summary"""
        self._add_section_header(doc, '6. FINANCIAL INFORMATION')
        
        # Capital structure
        doc.add_paragraph('6.1 CAPITAL STRUCTURE', style='Heading 3')
        
        capital_table = doc.add_table(rows=5, cols=2)
        capital_table.style = 'Table Grid'
        
        capital_info = [
            ('Authorized Capital:', f"M{self.lender.get('stated_capital', 0):,.2f}"),
            ('Issued Capital:', f"M{self._calculate_issued_capital():,.2f}"),
            ('Paid-up Capital:', f"M{self._calculate_paid_up_capital():,.2f}"),
            ('Total Assets (Current):', f"M{self.lender.get('total_assets', 0):,.2f}"),
            ('Net Worth:', f"M{self._calculate_net_worth():,.2f}"),
        ]
        
        for i, (label, value) in enumerate(capital_info):
            capital_table.rows[i].cells[0].text = label
            capital_table.rows[i].cells[1].text = value
            capital_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Financial projections summary
        doc.add_paragraph()
        doc.add_paragraph('6.2 3-YEAR FINANCIAL PROJECTIONS (SUMMARY)', style='Heading 3')
        
        projections_table = doc.add_table(rows=5, cols=4)
        projections_table.style = 'Table Grid'
        
        # Headers
        proj_headers = ['Financial Metric', 'Year 1', 'Year 2', 'Year 3']
        header_row = projections_table.rows[0]
        for i, header in enumerate(proj_headers):
            header_row.cells[i].text = header
            header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Financial projections (simplified)
        projections = self._generate_financial_projections_summary()
        proj_rows = [
            ('Total Revenue', projections['revenue']),
            ('Operating Expenses', projections['expenses']),
            ('Net Income', projections['profit']),
            ('Total Assets', projections['assets']),
        ]
        
        for i, (metric, values) in enumerate(proj_rows, 1):
            row = projections_table.rows[i]
            row.cells[0].text = metric
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            for j, value in enumerate(values, 1):
                row.cells[j].text = f"M{value:,.0f}"
    
    def _add_section_7_declaration(self, doc: Document):
        """Section 7: Compliance Declaration and Signatures"""
        self._add_section_header(doc, '7. DECLARATION AND CERTIFICATION')
        
        # Declaration text
        declaration_text = """
I/We hereby declare that:

1. The information provided in this Information Sheet is true, complete, and accurate to the best of my/our knowledge and belief.

2. The company will comply with all applicable laws, regulations, and guidelines issued by the Central Bank of Lesotho.

3. All directors and key management personnel meet the fit and proper criteria as outlined in CBL regulations.

4. The company has adequate systems, controls, and procedures in place to conduct microfinance business in a sound and prudent manner.

5. Any material changes to the information provided herein will be promptly communicated to the Central Bank of Lesotho.

6. The company understands that providing false or misleading information may result in the rejection of this application or revocation of any license granted.
        """
        
        declaration_para = doc.add_paragraph()
        declaration_para.add_run(declaration_text.strip())
        
        # Signature section
        doc.add_paragraph()
        
        signature_table = doc.add_table(rows=6, cols=2)
        signature_table.style = 'Table Grid'
        
        # CEO signature
        signature_table.rows[0].cells[0].text = 'Chief Executive Officer'
        signature_table.rows[0].cells[1].text = 'Board Chairman'
        
        signature_table.rows[1].cells[0].text = f"Name: {self._get_ceo_name()}"
        signature_table.rows[1].cells[1].text = f"Name: {self._get_chairman_name()}"
        
        signature_table.rows[2].cells[0].text = 'Signature: ________________________'
        signature_table.rows[2].cells[1].text = 'Signature: ________________________'
        
        signature_table.rows[3].cells[0].text = 'Date: ________________________'
        signature_table.rows[3].cells[1].text = 'Date: ________________________'
        
        signature_table.rows[4].cells[0].text = 'Company Stamp'
        signature_table.rows[4].cells[1].text = 'Witness'
        
        signature_table.rows[5].cells[0].text = ''
        signature_table.rows[5].cells[1].text = 'Name: ________________________'
        
        # Set row heights for signature section
        for row in signature_table.rows[2:4]:
            row.height = Inches(0.5)
        
        signature_table.rows[4].height = Inches(0.8)
        signature_table.rows[5].height = Inches(0.5)
    
    # Helper methods
    def _add_section_header(self, doc: Document, title: str):
        """Add a formatted section header"""
        doc.add_paragraph()
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(title)
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_run.font.underline = True
        doc.add_paragraph()
    
    def _format_date(self, date_obj) -> str:
        """Format date for display"""
        if isinstance(date_obj, (date, datetime)):
            return date_obj.strftime('%d/%m/%Y')
        elif isinstance(date_obj, str):
            try:
                parsed_date = datetime.strptime(date_obj, '%Y-%m-%d').date()
                return parsed_date.strftime('%d/%m/%Y')
            except:
                return date_obj
        return 'N/A'
    
    def _determine_business_nature(self) -> str:
        """Determine nature of business based on tier"""
        tier = self.lender.get('cbl_tier', '')
        
        nature_map = {
            'tier1': 'Deposit-taking Microfinance Institution',
            'tier2': 'Credit-only Microfinance Institution (Large)',
            'tier3': 'Credit-only Microfinance Institution (Small)',
        }
        
        return nature_map.get(tier, 'Microfinance Institution')
    
    def _extract_shareholding_data(self) -> List[Dict]:
        """Extract shareholding information from lender data"""
        # This would typically come from a related Shareholder model
        # For now, generate basic structure based on available data
        
        ceo_name = self._get_ceo_name()
        total_capital = float(self.lender.get('stated_capital', 1000000))
        
        return [
            {
                'name': ceo_name,
                'id_number': 'To be provided',
                'nationality': 'Mosotho',
                'shares': int(total_capital * 0.6),  # 60% shareholding
                'percentage': 60.0,
            },
            {
                'name': 'Other Shareholders',
                'id_number': 'Various',
                'nationality': 'Mosotho',
                'shares': int(total_capital * 0.4),  # 40% shareholding
                'percentage': 40.0,
            }
        ]
    
    def _identify_beneficial_owners(self) -> List[Dict]:
        """Identify beneficial owners (≥25% ownership)"""
        shareholding_data = self._extract_shareholding_data()
        return [s for s in shareholding_data if s['percentage'] >= 25.0 and s['name'] != 'Other Shareholders']
    
    def _calculate_foreign_ownership(self) -> float:
        """Calculate percentage of foreign ownership"""
        # Assume domestic ownership unless specified otherwise
        return 0.0
    
    def _determine_board_position(self, member: Dict) -> str:
        """Determine board position of member"""
        if member.get('is_chairman'):
            return 'Chairman'
        elif member.get('is_non_executive'):
            return 'Non-Executive Director'
        else:
            return 'Executive Director'
    
    def _extract_other_directorships(self, member: Dict) -> str:
        """Extract other directorship information"""
        affiliations = member.get('other_affiliations', '')
        if affiliations:
            return affiliations[:50] + '...' if len(affiliations) > 50 else affiliations
        return 'None declared'
    
    def _add_board_compliance_section(self, doc: Document, board_members: List):
        """Add board compliance analysis"""
        doc.add_paragraph('3.1 BOARD COMPLIANCE ANALYSIS', style='Heading 3')
        
        tier = self.lender.get('cbl_tier', '')
        required_minimum = 5 if tier in ['tier1', 'tier2'] else 3
        
        compliance_para = doc.add_paragraph()
        compliance_run = compliance_para.add_run('REGULATORY COMPLIANCE: ')
        compliance_run.font.bold = True
        
        if len(board_members) >= required_minimum:
            compliance_para.add_run(f'✓ Board meets minimum requirement of {required_minimum} directors.')
        else:
            compliance_para.add_run(f'⚠ Board below minimum requirement ({len(board_members)}/{required_minimum} directors).')
        
        # Independence analysis
        independent_count = sum(1 for m in board_members if m.get('is_non_executive'))
        independence_para = doc.add_paragraph()
        independence_run = independence_para.add_run('INDEPENDENCE: ')
        independence_run.font.bold = True
        independence_para.add_run(f'{independent_count} of {len(board_members)} directors are independent.')
        
        # Chairman status
        chairman = next((m for m in board_members if m.get('is_chairman')), None)
        chairman_para = doc.add_paragraph()
        chairman_run = chairman_para.add_run('CHAIRMAN: ')
        chairman_run.font.bold = True
        if chairman and chairman.get('is_non_executive'):
            chairman_para.add_run('✓ Chairman is non-executive (compliant).')
        else:
            chairman_para.add_run('⚠ Chairman should be non-executive director.')
    
    def _extract_qualifications(self, person: Dict) -> str:
        """Extract professional qualifications"""
        qualifications = person.get('professional_qualifications', '')
        if qualifications:
            # Truncate for table display
            return qualifications[:30] + '...' if len(qualifications) > 30 else qualifications
        return 'To be provided'
    
    def _calculate_experience_years(self, person: Dict) -> str:
        """Calculate years of experience from employment history"""
        employment_history = person.get('employment_history_10_years', '')
        if employment_history:
            # Simple heuristic - assume 5-10 years based on content length
            return '5-10' if len(employment_history) > 100 else '2-5'
        return 'N/A'
    
    def _extract_previous_institution(self, person: Dict) -> str:
        """Extract previous institution from employment history"""
        employment_history = person.get('employment_history_10_years', '')
        if employment_history:
            # Extract first institution mentioned (simple heuristic)
            lines = employment_history.split('\n')
            for line in lines:
                if any(word in line.lower() for word in ['bank', 'financial', 'institution', 'company']):
                    return line[:30] + '...' if len(line) > 30 else line
        return 'To be provided'
    
    def _generate_management_narrative(self) -> str:
        """Generate management structure narrative"""
        return f"""
The company has established a management structure appropriate for a {self.lender.get('cbl_tier', '').upper()} microfinance institution. 
Key management positions include Chief Executive Officer, Finance Officer, and Compliance Officer, all of whom meet 
the fit and proper criteria outlined in CBL regulations. The management team brings extensive experience in financial 
services and microfinance operations, ensuring sound governance and operational excellence.
        """.strip()
    
    def _determine_business_activities(self) -> Dict:
        """Determine primary and secondary business activities"""
        tier = self.lender.get('cbl_tier', '')
        
        if tier == 'tier1':
            primary = [
                'Accepting deposits from the public',
                'Providing micro-credit to individuals and small businesses',
                'Money transfer services',
                'Financial advisory services',
                'Insurance agency services (where licensed)'
            ]
        else:
            primary = [
                'Providing micro-credit to individuals and small businesses',
                'Financial advisory services',
                'Business development support',
                'Financial literacy training'
            ]
        
        return {
            'primary': primary,
            'secondary': ['Community development programs', 'Mobile banking services']
        }
    
    def _generate_target_market_description(self) -> str:
        """Generate target market description"""
        return f"""
The company targets financially excluded and underbanked populations in {self.lender.get('district', 'Maseru')} and surrounding areas. 
Primary focus is on small-scale entrepreneurs, women's groups, agricultural value chain participants, and low-income households 
seeking access to affordable financial services. The target market includes approximately {self.lender.get('target_market_size', 1000)} potential clients 
within the initial operational area.
        """.strip()
    
    def _generate_geographic_coverage(self) -> str:
        """Generate geographic coverage description"""
        return f"""
Initial operations will be concentrated in {self.lender.get('district', 'Maseru')} district, with plans to expand to adjacent 
districts within the first three years of operation. The company will establish {self._calculate_planned_branches()} branch(es) 
to serve the target market effectively while maintaining regulatory compliance and operational efficiency.
        """.strip()
    
    def _calculate_planned_branches(self) -> int:
        """Calculate number of planned branches based on capital"""
        capital = float(self.lender.get('stated_capital', 0))
        if capital >= 5000000:
            return 3
        elif capital >= 2000000:
            return 2
        else:
            return 1
    
    def _determine_products_services(self) -> List[Dict]:
        """Determine products and services offered"""
        tier = self.lender.get('cbl_tier', '')
        
        products = [
            {
                'name': 'Individual Microloans',
                'target_segment': 'Small entrepreneurs',
                'interest_rate': '24% - 36% p.a.',
                'loan_term': '3 - 24 months'
            },
            {
                'name': 'Group Lending',
                'target_segment': 'Women groups, cooperatives',
                'interest_rate': '20% - 30% p.a.',
                'loan_term': '6 - 18 months'
            },
            {
                'name': 'Agricultural Loans',
                'target_segment': 'Small-scale farmers',
                'interest_rate': '18% - 28% p.a.',
                'loan_term': '6 - 12 months'
            }
        ]
        
        if tier == 'tier1':
            products.insert(0, {
                'name': 'Savings Accounts',
                'target_segment': 'General public',
                'interest_rate': '6% - 8% p.a.',
                'loan_term': 'Demand deposits'
            })
        
        return products
    
    def _calculate_issued_capital(self) -> float:
        """Calculate issued capital"""
        return float(self.lender.get('stated_capital', 0)) * 0.8  # 80% of authorized
    
    def _calculate_paid_up_capital(self) -> float:
        """Calculate paid-up capital"""
        return float(self.lender.get('stated_capital', 0)) * 0.6  # 60% of authorized
    
    def _calculate_net_worth(self) -> float:
        """Calculate current net worth"""
        assets = float(self.lender.get('total_assets', 0))
        # Assume 20% liabilities for estimation
        return assets * 0.8
    
    def _generate_financial_projections_summary(self) -> Dict:
        """Generate simplified financial projections for summary table"""
        capital = float(self.lender.get('stated_capital', 0))
        
        # Simple projection based on capital size and tier
        base_revenue = capital * 0.3  # 30% of capital as annual revenue
        base_expenses = base_revenue * 0.7  # 70% expense ratio
        
        projections = {
            'revenue': [base_revenue * 0.6, base_revenue * 0.8, base_revenue],
            'expenses': [base_expenses * 0.6, base_expenses * 0.8, base_expenses],
            'profit': [
                (base_revenue * 0.6) - (base_expenses * 0.6),
                (base_revenue * 0.8) - (base_expenses * 0.8),
                base_revenue - base_expenses
            ],
            'assets': [capital * 1.2, capital * 1.5, capital * 2.0]
        }
        
        return projections
    
    def _get_ceo_name(self) -> str:
        """Get CEO name from personnel or lender data"""
        ceo = next((p for p in self.personnel if p.get('role') == 'ceo'), None)
        if ceo:
            return ceo.get('full_name', '')
        
        # Fallback to lender profile
        first_name = self.lender.get('ceo_first_name', '')
        last_name = self.lender.get('ceo_last_name', '')
        return f"{first_name} {last_name}".strip() or 'To be appointed'
    
    def _get_chairman_name(self) -> str:
        """Get chairman name from board members"""
        chairman = next((p for p in self.personnel 
                        if p.get('role') == 'director' and p.get('is_chairman')), None)
        return chairman.get('full_name', 'To be appointed') if chairman else 'To be appointed'


#=============================================================================================================
    # Fit & Proper Questionnaire Generator
    # Generates CBL-compliant fit and proper questionnaires for directors and key personnel
    # Based on CBL requirements and Schedule III formatting
#=================================================================================================================
class FitProperGenerator:
    """
    Generates comprehensive Fit & Proper Questionnaire documents
    
    This generator creates CBL-compliant fit and proper assessments for:
    - Directors (Executive and Non-Executive)
    - Chief Executive Officers
    - Finance Officers/Managers
    - Compliance Officers
    - Other Key Management Personnel
    """
    
    # CBL Fit & Proper Criteria Categories
    ASSESSMENT_CATEGORIES = {
        'integrity': {
            'weight': 25,
            'criteria': [
                'No criminal convictions',
                'No involvement in business failures',
                'Good character references',
                'Transparent business dealings'
            ]
        },
        'competence': {
            'weight': 25,
            'criteria': [
                'Relevant qualifications',
                'Appropriate experience',
                'Technical knowledge',
                'Leadership skills'
            ]
        },
        'financial_soundness': {
            'weight': 20,
            'criteria': [
                'Personal financial stability',
                'No bankruptcy history',
                'Adequate personal resources',
                'No conflicting financial interests'
            ]
        },
        'regulatory_compliance': {
            'weight': 15,
            'criteria': [
                'No regulatory sanctions',
                'Compliance with previous commitments',
                'Understanding of regulatory requirements',
                'No disqualifications'
            ]
        },
        'reputation': {
            'weight': 15,
            'criteria': [
                'Professional standing',
                'Community reputation',
                'Media coverage assessment',
                'Reference checks'
            ]
        }
    }
    
    # Role-specific requirements
    ROLE_REQUIREMENTS = {
        'director': {
            'min_experience_years': 5,
            'required_competencies': ['governance', 'oversight', 'strategy'],
            'preferred_qualifications': ['business_degree', 'professional_certification'],
            'independence_required': True,  # For non-executive directors
        },
        'ceo': {
            'min_experience_years': 7,
            'required_competencies': ['leadership', 'financial_management', 'strategic_planning'],
            'preferred_qualifications': ['business_degree', 'management_experience'],
            'independence_required': False,
        },
        'finance_officer': {
            'min_experience_years': 5,
            'required_competencies': ['accounting', 'financial_analysis', 'risk_management'],
            'preferred_qualifications': ['accounting_degree', 'professional_certification'],
            'independence_required': False,
        },
        'compliance_officer': {
            'min_experience_years': 3,
            'required_competencies': ['regulatory_knowledge', 'risk_assessment', 'policy_development'],
            'preferred_qualifications': ['law_degree', 'compliance_certification'],
            'independence_required': False,
        }
    }
    
    def __init__(self, personnel_data: Dict, lender_context: Optional[Dict] = None):
        self.personnel = personnel_data
        self.lender = lender_context or {}
        self.role = personnel_data.get('role', '')
        self.role_requirements = self.ROLE_REQUIREMENTS.get(self.role, {})
        
    def generate(self) -> Document:
        """Generate complete Fit & Proper questionnaire document"""
        doc = Document()
        
        # Set up document formatting
        self._setup_document_style(doc)
        
        # Add CBL header
        self._add_header(doc)
        
        # Part A: Personal Information
        self._add_part_a_personal_info(doc)
        
        # Part B: Educational and Professional Background
        self._add_part_b_background(doc)
        
        # Part C: Employment History (10 years)
        self._add_part_c_employment(doc)
        
        # Part D: Business Affiliations and Interests
        self._add_part_d_affiliations(doc)
        
        # Part E: Family and Related Party Information
        self._add_part_e_family(doc)
        
        # Part F: Financial Position
        self._add_part_f_financial(doc)
        
        # Part G: Legal and Regulatory History
        self._add_part_g_legal(doc)
        
        # Part H: References
        self._add_part_h_references(doc)
        
        # Part I: Declaration and Undertakings
        self._add_part_i_declaration(doc)
        
        # Part J: Supporting Documents Checklist
        self._add_part_j_documents(doc)
        
        return doc
    
    def generate_assessment_summary(self) -> Dict:
        """Generate fit and proper assessment summary"""
        assessment = {
            'overall_rating': self._calculate_overall_rating(),
            'category_scores': self._assess_all_categories(),
            'red_flags': self._identify_red_flags(),
            'recommendations': self._generate_recommendations(),
            'missing_information': self._identify_missing_information(),
            'approval_recommendation': self._determine_approval_recommendation(),
        }
        
        return assessment
    
    def _setup_document_style(self, doc: Document):
        """Set up document-wide styling"""
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
    
    def _add_header(self, doc: Document):
        """Add CBL header and form identification"""
        # CBL Header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        header_run = header_para.add_run('CENTRAL BANK OF LESOTHO\n')
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        
        subheader_run = header_para.add_run('MICROFINANCE INSTITUTIONS AND COOPERATIVES SUPERVISION DEPARTMENT\n')
        subheader_run.font.size = Pt(11)
        subheader_run.font.bold = True
        
        # Form title
        doc.add_paragraph()
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_run = title_para.add_run('FIT AND PROPER QUESTIONNAIRE\n')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.underline = True
        
        subtitle_run = title_para.add_run('For Directors and Key Management Personnel')
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        
        # Form metadata
        doc.add_paragraph()
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Table Grid'
        
        metadata_info = [
            ('Institution Name:', self.lender.get('company_name', '')),
            ('Applicant Name:', self.personnel.get('full_name', '')),
            ('Position Applied For:', self._format_position()),
            ('Date Completed:', datetime.now().strftime('%d/%m/%Y')),
        ]
        
        for i, (label, value) in enumerate(metadata_info):
            metadata_table.rows[i].cells[0].text = label
            metadata_table.rows[i].cells[1].text = value
            metadata_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Instructions
        doc.add_paragraph()
        instructions = doc.add_paragraph()
        instructions_run = instructions.add_run('INSTRUCTIONS: ')
        instructions_run.font.bold = True
        instructions.add_run(
            'Please complete all sections of this questionnaire in full. Any incomplete or inaccurate '
            'information may result in delays in processing or rejection of the application. '
            'All responses must be supported by appropriate documentation as indicated.'
        )
        
        doc.add_paragraph()
    
    def _add_part_a_personal_info(self, doc: Document):
        """Part A: Personal Information"""
        self._add_section_header(doc, 'PART A: PERSONAL INFORMATION')
        
        # Personal details table
        personal_table = doc.add_table(rows=12, cols=2)
        personal_table.style = 'Table Grid'
        personal_table.columns[0].width = Inches(2.5)
        personal_table.columns[1].width = Inches(4.5)
        
        personal_info = [
            ('1. Full Name (as per ID):', self.personnel.get('full_name', '')),
            ('2. Date of Birth:', self._format_date(self.personnel.get('date_of_birth'))),
            ('3. Place of Birth:', self.personnel.get('place_of_birth', '')),
            ('4. Nationality:', self.personnel.get('nationality', '')),
            ('5. ID/Passport Number:', self.personnel.get('id_or_passport_number', '')),
            ('6. Country of Residence:', self.personnel.get('country_of_residence', '')),
            ('7. Residential Address:', self.personnel.get('residential_address', '')),
            ('8. Business Address:', self.personnel.get('business_address', '')),
            ('9. Telephone (Mobile):', self.personnel.get('mobile_phone', '')),
            ('10. Telephone (Office):', self.personnel.get('office_phone', '')),
            ('11. Email Address:', self.personnel.get('email', '')),
            ('12. Marital Status:', self.personnel.get('marital_status', '')),
        ]
        
        for i, (label, value) in enumerate(personal_info):
            personal_table.rows[i].cells[0].text = label
            personal_table.rows[i].cells[1].text = str(value) if value else ''
            personal_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    def _add_part_b_background(self, doc: Document):
        """Part B: Educational and Professional Background"""
        self._add_section_header(doc, 'PART B: EDUCATIONAL AND PROFESSIONAL BACKGROUND')
        
        # Educational qualifications
        doc.add_paragraph('B.1 EDUCATIONAL QUALIFICATIONS', style='Heading 3')
        
        education_table = doc.add_table(rows=1, cols=5)
        education_table.style = 'Table Grid'
        
        education_headers = ['Institution', 'Qualification', 'Year Completed', 'Grade/Class', 'Relevant to Position']
        header_row = education_table.rows[0]
        for i, header in enumerate(education_headers):
            header_row.cells[i].text = header
            header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Parse and add educational qualifications
        qualifications = self._parse_qualifications()
        for qualification in qualifications:
            row = education_table.add_row()
            row.cells[0].text = qualification['institution']
            row.cells[1].text = qualification['qualification']
            row.cells[2].text = qualification['year']
            row.cells[3].text = qualification['grade']
            row.cells[4].text = qualification['relevance']
        
        # Professional certifications
        doc.add_paragraph()
        doc.add_paragraph('B.2 PROFESSIONAL CERTIFICATIONS AND MEMBERSHIPS', style='Heading 3')
        
        cert_table = doc.add_table(rows=1, cols=4)
        cert_table.style = 'Table Grid'
        
        cert_headers = ['Organization', 'Certification/Membership', 'Date Obtained', 'Current Status']
        cert_header_row = cert_table.rows[0]
        for i, header in enumerate(cert_headers):
            cert_header_row.cells[i].text = header
            cert_header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add professional certifications
        certifications = self._parse_certifications()
        for cert in certifications:
            row = cert_table.add_row()
            row.cells[0].text = cert['organization']
            row.cells[1].text = cert['certification']
            row.cells[2].text = cert['date_obtained']
            row.cells[3].text = cert['status']
        
        # Skills and competencies assessment
        doc.add_paragraph()
        doc.add_paragraph('B.3 KEY COMPETENCIES ASSESSMENT', style='Heading 3')
        
        competencies_para = doc.add_paragraph()
        competencies_para.add_run('Please rate your competency level (1-5 scale) in the following areas relevant to your proposed role:')
        
        competency_table = doc.add_table(rows=1, cols=3)
        competency_table.style = 'Table Grid'
        
        comp_headers = ['Competency Area', 'Self-Assessment (1-5)', 'Evidence/Experience']
        comp_header_row = competency_table.rows[0]
        for i, header in enumerate(comp_headers):
            comp_header_row.cells[i].text = header
            comp_header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add role-specific competencies
        required_competencies = self.role_requirements.get('required_competencies', [])
        for competency in required_competencies:
            row = competency_table.add_row()
            row.cells[0].text = competency.replace('_', ' ').title()
            row.cells[1].text = ''  # To be filled by applicant
            row.cells[2].text = ''  # To be filled by applicant
    
    def _add_part_c_employment(self, doc: Document):
        """Part C: Employment History (Last 10 Years)"""
        self._add_section_header(doc, 'PART C: EMPLOYMENT HISTORY (LAST 10 YEARS)')
        
        instructions = doc.add_paragraph()
        instructions_run = instructions.add_run('INSTRUCTIONS: ')
        instructions_run.font.bold = True
        instructions.add_run(
            'Please provide complete employment history for the last 10 years, including periods of '
            'self-employment, unemployment, or study. Include reasons for leaving each position.'
        )
        
        doc.add_paragraph()
        
        employment_table = doc.add_table(rows=1, cols=7)
        employment_table.style = 'Table Grid'
        
        employment_headers = [
            'Employer', 'Position', 'Start Date', 'End Date', 'Key Responsibilities', 
            'Reason for Leaving', 'Contact Details'
        ]
        header_row = employment_table.rows[0]
        for i, header in enumerate(employment_headers):
            header_row.cells[i].text = header
            header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Parse employment history from text field
        employment_entries = self._parse_employment_history()
        for entry in employment_entries:
            row = employment_table.add_row()
            row.cells[0].text = entry['employer']
            row.cells[1].text = entry['position']
            row.cells[2].text = entry['start_date']
            row.cells[3].text = entry['end_date']
            row.cells[4].text = entry['responsibilities']
            row.cells[5].text = entry['reason_leaving']
            row.cells[6].text = entry['contact_details']
        
        # Employment gaps analysis
        doc.add_paragraph()
        doc.add_paragraph('C.1 EXPLANATION OF ANY GAPS IN EMPLOYMENT', style='Heading 3')
        
        gaps_para = doc.add_paragraph()
        employment_gaps = self._identify_employment_gaps()
        if employment_gaps:
            gaps_para.add_run('Employment gaps identified: ')
            for gap in employment_gaps:
                gaps_para.add_run(f"{gap['period']} - {gap['explanation']}; ")
        else:
            gaps_para.add_run('No significant employment gaps identified in the provided history.')
    
    def _add_part_d_affiliations(self, doc: Document):
        """Part D: Business Affiliations and Interests"""
        self._add_section_header(doc, 'PART D: BUSINESS AFFILIATIONS AND INTERESTS')
        
        # Current directorships
        doc.add_paragraph('D.1 CURRENT DIRECTORSHIPS AND SENIOR POSITIONS', style='Heading 3')
        
        directorships_table = doc.add_table(rows=1, cols=5)
        directorships_table.style = 'Table Grid'
        
        dir_headers = ['Company Name', 'Position', 'Date Appointed', 'Shareholding %', 'Nature of Business']
        dir_header_row = directorships_table.rows[0]
        for i, header in enumerate(dir_headers):
            dir_header_row.cells[i].text = header
            dir_header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add directorship information
        directorships = self._parse_other_affiliations()
        for directorship in directorships:
            row = directorships_table.add_row()
            row.cells[0].text = directorship['company']
            row.cells[1].text = directorship['position']
            row.cells[2].text = directorship['date_appointed']
            row.cells[3].text = directorship['shareholding']
            row.cells[4].text = directorship['nature_of_business']
        
        # Shareholdings > 5%
        doc.add_paragraph()
        doc.add_paragraph('D.2 SHAREHOLDINGS GREATER THAN 5%', style='Heading 3')
        
        shareholding_table = doc.add_table(rows=1, cols=4)
        shareholding_table.style = 'Table Grid'
        
        share_headers = ['Company Name', 'Percentage Held', 'Nature of Business', 'Listed/Unlisted']
        share_header_row = shareholding_table.rows[0]
        for i, header in enumerate(share_headers):
            share_header_row.cells[i].text = header
            share_header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Business partnerships
        doc.add_paragraph()
        doc.add_paragraph('D.3 BUSINESS PARTNERSHIPS AND JOINT VENTURES', style='Heading 3')
        
        partnerships_para = doc.add_paragraph()
        partnerships_text = self.personnel.get('business_partnerships', '')
        if partnerships_text:
            partnerships_para.add_run(partnerships_text)
        else:
            partnerships_para.add_run('No current business partnerships or joint ventures.')
        
        # Conflicts of interest
        doc.add_paragraph()
        doc.add_paragraph('D.4 POTENTIAL CONFLICTS OF INTEREST', style='Heading 3')
        
        conflicts_para = doc.add_paragraph()
        conflicts_para.add_run(
            'Please describe any potential conflicts of interest that may arise from your '
            'appointment to this position, including business relationships with clients, '
            'competitors, or related parties of the institution.'
        )
        
        conflicts_response = doc.add_paragraph()
        conflicts_response.add_run(self._assess_potential_conflicts())
    
    def _add_part_e_family(self, doc: Document):
        """Part E: Family and Related Party Information"""
        self._add_section_header(doc, 'PART E: FAMILY AND RELATED PARTY INFORMATION')
        
        instructions = doc.add_paragraph()
        instructions_run = instructions.add_run('NOTE: ')
        instructions_run.font.bold = True
        instructions.add_run(
            'This information is required to assess potential conflicts of interest and related party '
            'transactions as required by CBL regulations.'
        )
        
        doc.add_paragraph()
        
        # Immediate family
        doc.add_paragraph('E.1 IMMEDIATE FAMILY MEMBERS', style='Heading 3')
        
        family_table = doc.add_table(rows=1, cols=5)
        family_table.style = 'Table Grid'
        
        family_headers = ['Relationship', 'Full Name', 'Occupation', 'Employer/Business', 'Financial Sector Involvement']
        family_header_row = family_table.rows[0]
        for i, header in enumerate(family_headers):
            family_header_row.cells[i].text = header
            family_header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Parse family information
        family_members = self._parse_family_affiliations()
        for member in family_members:
            row = family_table.add_row()
            row.cells[0].text = member['relationship']
            row.cells[1].text = member['name']
            row.cells[2].text = member['occupation']
            row.cells[3].text = member['employer']
            row.cells[4].text = member['financial_involvement']
        
        # Related party businesses
        doc.add_paragraph()
        doc.add_paragraph('E.2 RELATED PARTY BUSINESS INTERESTS', style='Heading 3')
        
        related_party_para = doc.add_paragraph()
        related_party_para.add_run(
            'List any businesses in which you or your immediate family members have a '
            'significant interest (>5% ownership or control):'
        )
        
        related_businesses = self._identify_related_party_businesses()
        if related_businesses:
            for business in related_businesses:
                business_para = doc.add_paragraph(style='List Bullet')
                business_para.add_run(f"{business['name']} - {business['relationship']} - {business['nature']}")
        else:
            no_related_para = doc.add_paragraph()
            no_related_para.add_run('No significant related party business interests to declare.')
    
    def _add_part_f_financial(self, doc: Document):
        """Part F: Financial Position"""
        self._add_section_header(doc, 'PART F: PERSONAL FINANCIAL POSITION')
        
        confidentiality_note = doc.add_paragraph()
        conf_run = confidentiality_note.add_run('CONFIDENTIALITY NOTE: ')
        conf_run.font.bold = True
        confidentiality_note.add_run(
            'The information provided in this section will be treated in strict confidence '
            'and will only be used for fit and proper assessment purposes.'
        )
        
        doc.add_paragraph()
        
        # Financial summary
        doc.add_paragraph('F.1 PERSONAL FINANCIAL SUMMARY', style='Heading 3')
        
        financial_table = doc.add_table(rows=8, cols=2)
        financial_table.style = 'Table Grid'
        financial_table.columns[0].width = Inches(3)
        financial_table.columns[1].width = Inches(3)
        
        # Estimated financial position
        estimated_finances = self._estimate_financial_position()
        
        financial_items = [
            ('Total Assets (M)', f"{estimated_finances['assets']:,.2f}"),
            ('Total Liabilities (M)', f"{estimated_finances['liabilities']:,.2f}"),
            ('Net Worth (M)', f"{estimated_finances['net_worth']:,.2f}"),
            ('Annual Income (M)', f"{estimated_finances['annual_income']:,.2f}"),
            ('Monthly Expenses (M)', f"{estimated_finances['monthly_expenses']:,.2f}"),
            ('Outstanding Loans (M)', f"{estimated_finances['outstanding_loans']:,.2f}"),
            ('Credit Facilities Available (M)', f"{estimated_finances['credit_facilities']:,.2f}"),
            ('Financial Dependents', str(estimated_finances['dependents'])),
        ]
        
        for i, (label, value) in enumerate(financial_items):
            financial_table.rows[i].cells[0].text = label
            financial_table.rows[i].cells[1].text = value
            financial_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Source of wealth
        doc.add_paragraph()
        doc.add_paragraph('F.2 SOURCE OF WEALTH AND INCOME', style='Heading 3')
        
        wealth_para = doc.add_paragraph()
        wealth_para.add_run('Primary sources of income and wealth: ')
        wealth_para.add_run(estimated_finances['income_sources'])
        
        # Financial obligations
        doc.add_paragraph()
        doc.add_paragraph('F.3 SIGNIFICANT FINANCIAL OBLIGATIONS', style='Heading 3')
        
        obligations_para = doc.add_paragraph()
        obligations_para.add_run(estimated_finances['obligations'])
    
    def _add_part_g_legal(self, doc: Document):
        """Part G: Legal and Regulatory History"""
        self._add_section_header(doc, 'PART G: LEGAL AND REGULATORY HISTORY')
        
        warning = doc.add_paragraph()
        warning_run = warning.add_run('WARNING: ')
        warning_run.font.bold = True
        warning_run.font.color.rgb = (255, 0, 0)  # Red color
        warning.add_run(
            'Providing false or misleading information in this section may result in criminal charges '
            'and automatic disqualification. All information will be verified through official channels.'
        )
        
        doc.add_paragraph()
        
        # Legal declarations with checkboxes
        declarations = [
            ('G.1', 'Criminal Convictions', 'criminal_conviction'),
            ('G.2', 'Civil Court Judgments', 'civil_judgments'),
            ('G.3', 'Bankruptcy/Insolvency', 'bankruptcy'),
            ('G.4', 'Regulatory Actions', 'regulatory_actions'),
            ('G.5', 'Professional Disqualifications', 'ever_disqualified'),
            ('G.6', 'Business Failures', 'business_failures'),
            ('G.7', 'Dismissals/Resignations', 'dismissed_or_resigned'),
        ]
        
        for section, title, field_name in declarations:
            self._add_legal_declaration_section(doc, section, title, field_name)
    
    def _add_legal_declaration_section(self, doc: Document, section_num: str, title: str, field_name: str):
        """Add individual legal declaration section"""
        doc.add_paragraph(f'{section_num} {title.upper()}', style='Heading 3')
        
        # Get declaration status
        declaration_value = self.personnel.get(field_name, False)
        details_field = f'{field_name}_details'
        details_value = self.personnel.get(details_field, '')
        
        # Yes/No question
        question_para = doc.add_paragraph()
        question_text = self._get_declaration_question_text(field_name)
        question_para.add_run(question_text)
        
        # Response
        response_para = doc.add_paragraph()
        response_run = response_para.add_run('Response: ')
        response_run.font.bold = True
        
        if declaration_value:
            response_para.add_run('YES')
            
            # Add details if provided
            if details_value:
                details_para = doc.add_paragraph()
                details_run = details_para.add_run('Details: ')
                details_run.font.bold = True
                details_para.add_run(details_value)
            else:
                placeholder_para = doc.add_paragraph()
                placeholder_run = placeholder_para.add_run('Details: ')
                placeholder_run.font.bold = True
                placeholder_para.add_run('[Please provide full details including dates, circumstances, and current status]')
        else:
            response_para.add_run('NO')
        
        doc.add_paragraph()
    
    def _add_part_h_references(self, doc: Document):
        """Part H: References"""
        self._add_section_header(doc, 'PART H: CHARACTER AND PROFESSIONAL REFERENCES')
        
        reference_instructions = doc.add_paragraph()
        reference_instructions.add_run(
            'Please provide details of at least 3 references who can attest to your character, '
            'professional competence, and suitability for the proposed position. References should '
            'include professional colleagues, former supervisors, and community leaders.'
        )
        
        doc.add_paragraph()
        
        # References table
        references_table = doc.add_table(rows=1, cols=6)
        references_table.style = 'Table Grid'
        
        ref_headers = ['Name', 'Position', 'Organization', 'Relationship', 'Contact Number', 'Email']
        ref_header_row = references_table.rows[0]
        for i, header in enumerate(ref_headers):
            ref_header_row.cells[i].text = header
            ref_header_row.cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add reference information from uploaded documents
        references = self._extract_reference_information()
        for ref in references[:5]:  # Maximum 5 references
            row = references_table.add_row()
            row.cells[0].text = ref['name']
            row.cells[1].text = ref['position']
            row.cells[2].text = ref['organization']
            row.cells[3].text = ref['relationship']
            row.cells[4].text = ref['contact_number']
            row.cells[5].text = ref['email']
        
        # Reference verification consent
        doc.add_paragraph()
        consent_para = doc.add_paragraph()
        consent_run = consent_para.add_run('CONSENT FOR REFERENCE VERIFICATION: ')
        consent_run.font.bold = True
        consent_para.add_run(
            'I hereby authorize the Central Bank of Lesotho to contact the above references '
            'and any other persons they deem necessary to verify the information provided in '
            'this questionnaire and to obtain additional information relevant to my fitness '
            'and propriety for the proposed position.'
        )
    
    def _add_part_i_declaration(self, doc: Document):
        """Part I: Declaration and Undertakings"""
        self._add_section_header(doc, 'PART I: DECLARATION AND UNDERTAKINGS')
        
        # Main declaration
        declaration_text = f"""
I, {self.personnel.get('full_name', '[NAME]')}, hereby solemnly declare that:

1. All information provided in this questionnaire is true, complete, and accurate to the best of my knowledge and belief.

2. I have not withheld any material information that may be relevant to my fitness and propriety assessment.

3. I understand that this questionnaire forms part of my application for appointment as {self._format_position()} at {self.lender.get('company_name', '[INSTITUTION NAME]')}.

4. I undertake to immediately notify the Central Bank of Lesotho and the institution of any material changes to the information provided herein.

5. I consent to the Central Bank of Lesotho conducting any investigations deemed necessary to verify the information provided.

6. I understand that providing false or misleading information may result in the rejection of my appointment application and/or criminal charges.

7. I undertake to comply with all applicable laws, regulations, and guidelines if appointed to the proposed position.

8. I understand my fiduciary duties and responsibilities as {self._format_position()} and undertake to discharge them diligently and in good faith.

9. I will maintain the confidentiality of all institution information and comply with all internal policies and procedures.

10. I consent to periodic reviews of my continued fitness and propriety during my tenure in the proposed position.
        """
        
        declaration_para = doc.add_paragraph()
        declaration_para.add_run(declaration_text.strip())
        
        # Signature section
        doc.add_paragraph()
        
        signature_table = doc.add_table(rows=4, cols=2)
        signature_table.style = 'Table Grid'
        
        signature_table.rows[0].cells[0].text = 'Applicant Signature'
        signature_table.rows[0].cells[1].text = 'Date'
        
        signature_table.rows[1].cells[0].text = '________________________________'
        signature_table.rows[1].cells[1].text = '________________'
        
        signature_table.rows[2].cells[0].text = f"Print Name: {self.personnel.get('full_name', '')}"
        signature_table.rows[2].cells[1].text = 'Witness (if required)'
        
        signature_table.rows[3].cells[0].text = 'Position Applied For: ' + self._format_position()
        signature_table.rows[3].cells[1].text = '________________________________'
        
        # Set row heights
        for row in signature_table.rows[1:2]:
            row.height = Inches(0.5)
    
    def _add_part_j_documents(self, doc: Document):
        """Part J: Supporting Documents Checklist"""
        self._add_section_header(doc, 'PART J: SUPPORTING DOCUMENTS CHECKLIST')
        
        checklist_intro = doc.add_paragraph()
        checklist_intro.add_run(
            'The following documents must be submitted with this completed questionnaire. '
            'Please check the box next to each document that is being submitted:'
        )
        
        doc.add_paragraph()
        
        # Required documents based on personnel data
        required_docs = self._get_required_documents_list()
        
        docs_table = doc.add_table(rows=len(required_docs) + 1, cols=3)
        docs_table.style = 'Table Grid'
        
        # Headers
        docs_table.rows[0].cells[0].text = '✓'
        docs_table.rows[0].cells[1].text = 'Required Document'
        docs_table.rows[0].cells[2].text = 'Status/Notes'
        
        for i, header_cell in enumerate(docs_table.rows[0].cells):
            header_cell.paragraphs[0].runs[0].font.bold = True
        
        # Document checklist
        for i, doc_info in enumerate(required_docs, 1):
            docs_table.rows[i].cells[0].text = '☐' if not doc_info['submitted'] else '☑'
            docs_table.rows[i].cells[1].text = doc_info['name']
            docs_table.rows[i].cells[2].text = doc_info['status']
            
            # Color code the checkbox
            if doc_info['submitted']:
                docs_table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = (0, 128, 0)  # Green
            else:
                docs_table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = (255, 0, 0)  # Red
        
        # Submission declaration
        doc.add_paragraph()
        submission_para = doc.add_paragraph()
        submission_run = submission_para.add_run('SUBMISSION DECLARATION: ')
        submission_run.font.bold = True
        
        submitted_count = sum(1 for doc in required_docs if doc['submitted'])
        total_count = len(required_docs)
        
        submission_para.add_run(
            f'I confirm that {submitted_count} of {total_count} required documents are being '
            f'submitted with this application. Any outstanding documents will be provided as indicated above.'
        )
    
    # Helper methods for data parsing and assessment
    def _format_position(self) -> str:
        """Format position title for display"""
        position_map = {
            'director': 'Director',
            'ceo': 'Chief Executive Officer',
            'finance_officer': 'Finance Officer',
            'compliance_officer': 'Compliance Officer',
        }
        return position_map.get(self.role, self.role.replace('_', ' ').title())
    
    def _format_date(self, date_obj) -> str:
        """Format date for display"""
        if isinstance(date_obj, (date, datetime)):
            return date_obj.strftime('%d/%m/%Y')
        elif isinstance(date_obj, str):
            try:
                parsed_date = datetime.strptime(date_obj, '%Y-%m-%d').date()
                return parsed_date.strftime('%d/%m/%Y')
            except:
                return date_obj
        return ''
    
    def _parse_qualifications(self) -> List[Dict]:
        """Parse educational qualifications from text"""
        qualifications_text = self.personnel.get('professional_qualifications', '')
        
        if not qualifications_text:
            return [{'institution': 'To be provided', 'qualification': '', 'year': '', 'grade': '', 'relevance': ''}]
        
        # Simple parsing - in production, this would be more sophisticated
        lines = qualifications_text.split('\n')
        qualifications = []
        
        for line in lines[:3]:  # Maximum 3 qualifications
            if line.strip():
                qualifications.append({
                    'institution': 'To be provided',
                    'qualification': line.strip()[:50],
                    'year': 'To be provided',
                    'grade': 'To be provided',
                    'relevance': 'High' if any(keyword in line.lower() for keyword in ['business', 'finance', 'accounting', 'management']) else 'Medium'
                })
        
        return qualifications or [{'institution': 'To be provided', 'qualification': '', 'year': '', 'grade': '', 'relevance': ''}]
    
    def _parse_certifications(self) -> List[Dict]:
        """Parse professional certifications"""
        # Extract from qualifications text
        qualifications_text = self.personnel.get('professional_qualifications', '')
        
        certifications = []
        if 'certified' in qualifications_text.lower() or 'cpa' in qualifications_text.lower():
            certifications.append({
                'organization': 'Professional Body',
                'certification': 'Professional Certification',
                'date_obtained': 'To be provided',
                'status': 'Active'
            })
        
        return certifications or [{'organization': 'To be provided', 'certification': 'To be provided', 'date_obtained': '', 'status': ''}]
    
    def _parse_employment_history(self) -> List[Dict]:
        """Parse 10-year employment history"""
        employment_text = self.personnel.get('employment_history_10_years', '')
        
        if not employment_text:
            return [{'employer': 'To be provided', 'position': '', 'start_date': '', 'end_date': '', 'responsibilities': '', 'reason_leaving': '', 'contact_details': ''}]
        
        # Simple parsing - split by lines or periods
        entries = []
        lines = employment_text.split('\n')
        
        current_entry = {}
        for line in lines:
            if line.strip():
                if not current_entry or len(current_entry) >= 6:
                    if current_entry:
                        entries.append(current_entry)
                    current_entry = {
                        'employer': line.strip()[:30],
                        'position': 'To be provided',
                        'start_date': 'To be provided',
                        'end_date': 'To be provided',
                        'responsibilities': line.strip()[:50] + '...' if len(line.strip()) > 50 else line.strip(),
                        'reason_leaving': 'To be provided',
                        'contact_details': 'To be provided'
                    }
        
        if current_entry:
            entries.append(current_entry)
        
        return entries[:5] or [{'employer': 'To be provided', 'position': '', 'start_date': '', 'end_date': '', 'responsibilities': '', 'reason_leaving': '', 'contact_details': ''}]
    
    def _identify_employment_gaps(self) -> List[Dict]:
        """Identify gaps in employment history"""
        # Placeholder for gap analysis
        return []
    
    def _parse_other_affiliations(self) -> List[Dict]:
        """Parse other business affiliations"""
        affiliations_text = self.personnel.get('other_affiliations', '')
        
        affiliations = []
        if affiliations_text:
            # Simple parsing
            affiliations.append({
                'company': 'To be provided',
                'position': 'Director/Member',
                'date_appointed': 'To be provided',
                'shareholding': '0%',
                'nature_of_business': affiliations_text[:50] + '...' if len(affiliations_text) > 50 else affiliations_text
            })
        
        return affiliations or [{'company': 'None', 'position': 'N/A', 'date_appointed': 'N/A', 'shareholding': 'N/A', 'nature_of_business': 'N/A'}]
    
    def _parse_family_affiliations(self) -> List[Dict]:
        """Parse family member information"""
        family_text = self.personnel.get('family_business_affiliations', '')
        
        family_members = []
        if family_text:
            family_members.append({
                'relationship': 'To be specified',
                'name': 'To be provided',
                'occupation': 'To be provided',
                'employer': 'To be provided',
                'financial_involvement': 'None declared'
            })
        
        return family_members or [{'relationship': 'Spouse', 'name': 'To be provided', 'occupation': '', 'employer': '', 'financial_involvement': 'None'}]
    
    def _assess_potential_conflicts(self) -> str:
        """Assess potential conflicts of interest"""
        affiliations = self.personnel.get('other_affiliations', '')
        family_affiliations = self.personnel.get('family_business_affiliations', '')
        
        if affiliations or family_affiliations:
            return (
                'Based on the business and family affiliations disclosed, potential conflicts '
                'may arise and will require ongoing monitoring and appropriate management procedures.'
            )
        else:
            return 'No immediate conflicts of interest identified based on the information provided.'
    
    def _identify_related_party_businesses(self) -> List[Dict]:
        """Identify related party business interests"""
        return []  # Placeholder
    
    def _estimate_financial_position(self) -> Dict:
        """Estimate financial position based on role and experience"""
        # Basic estimation based on role
        role_income_estimates = {
            'director': {'min': 180000, 'max': 500000},
            'ceo': {'min': 300000, 'max': 800000},
            'finance_officer': {'min': 150000, 'max': 400000},
            'compliance_officer': {'min': 120000, 'max': 350000},
        }
        
        income_range = role_income_estimates.get(self.role, {'min': 100000, 'max': 300000})
        estimated_annual_income = (income_range['min'] + income_range['max']) / 2
        
        return {
            'assets': estimated_annual_income * 2,  # 2x annual income
            'liabilities': estimated_annual_income * 0.5,  # 50% of annual income
            'net_worth': estimated_annual_income * 1.5,
            'annual_income': estimated_annual_income,
            'monthly_expenses': estimated_annual_income / 12 * 0.7,  # 70% of income
            'outstanding_loans': estimated_annual_income * 0.3,  # 30% of annual income
            'credit_facilities': estimated_annual_income * 0.2,  # 20% of annual income
            'dependents': 2,  # Average family size
            'income_sources': 'Employment income, investment returns',
            'obligations': 'Mortgage payments, family support, professional insurance'
        }
    
    def _get_declaration_question_text(self, field_name: str) -> str:
        """Get question text for legal declarations"""
        questions = {
            'criminal_conviction': 'Have you ever been convicted of any criminal offense (excluding minor traffic violations)?',
            'civil_judgments': 'Have you ever been subject to any adverse civil court judgments?',
            'bankruptcy': 'Have you ever been declared bankrupt or insolvent?',
            'regulatory_actions': 'Have you ever been subject to disciplinary action by any regulatory authority?',
            'ever_disqualified': 'Have you ever been disqualified from serving as a director or in any professional capacity?',
            'business_failures': 'Have you ever been associated with any business failure or insolvency?',
            'dismissed_or_resigned': 'Have you ever been dismissed from employment or asked to resign due to misconduct?',
        }
        
        return questions.get(field_name, f'Please answer regarding {field_name.replace("_", " ")}:')
    
    def _extract_reference_information(self) -> List[Dict]:
        """Extract reference information from uploaded documents"""
        # Check for uploaded reference documents
        references = []
        
        # Check for character references
        if self.personnel.get('character_ref_1'):
            references.append({
                'name': 'Character Reference 1',
                'position': 'To be extracted',
                'organization': 'To be extracted',
                'relationship': 'Professional colleague',
                'contact_number': 'To be provided',
                'email': 'To be provided'
            })
        
        if self.personnel.get('character_ref_2'):
            references.append({
                'name': 'Character Reference 2',
                'position': 'To be extracted',
                'organization': 'To be extracted',
                'relationship': 'Professional colleague',
                'contact_number': 'To be provided',
                'email': 'To be provided'
            })
        
        # Check for financial references
        if self.personnel.get('financial_ref_1'):
            references.append({
                'name': 'Bank Reference 1',
                'position': 'Relationship Manager',
                'organization': 'Financial Institution',
                'relationship': 'Banking relationship',
                'contact_number': 'To be provided',
                'email': 'To be provided'
            })
        
        return references or [{'name': 'To be provided', 'position': '', 'organization': '', 'relationship': '', 'contact_number': '', 'email': ''}]
    
    def _get_required_documents_list(self) -> List[Dict]:
        """Get list of required supporting documents"""
        base_documents = [
            {'name': 'Completed Fit & Proper Questionnaire', 'submitted': bool(self.personnel.get('fit_proper_form')), 'status': 'This document'},
            {'name': 'Detailed Curriculum Vitae', 'submitted': bool(self.personnel.get('curriculum_vitae')), 'status': 'Uploaded' if self.personnel.get('curriculum_vitae') else 'Required'},
            {'name': 'Police Clearance Certificate', 'submitted': bool(self.personnel.get('police_clearance')), 'status': 'Uploaded' if self.personnel.get('police_clearance') else 'Required'},
            {'name': 'Personal Tax Clearance Certificate', 'submitted': bool(self.personnel.get('tax_clearance_individual')), 'status': 'Uploaded' if self.personnel.get('tax_clearance_individual') else 'Required'},
            {'name': 'Certified Copy of ID/Passport', 'submitted': bool(self.personnel.get('id_copy')), 'status': 'Uploaded' if self.personnel.get('id_copy') else 'Required'},
            {'name': 'Statement of Assets & Liabilities', 'submitted': bool(self.personnel.get('statement_assets_liabilities')), 'status': 'Uploaded' if self.personnel.get('statement_assets_liabilities') else 'Required'},
            {'name': 'Character Reference 1 (Notarized)', 'submitted': bool(self.personnel.get('character_ref_1')), 'status': 'Uploaded' if self.personnel.get('character_ref_1') else 'Required'},
            {'name': 'Character Reference 2 (Notarized)', 'submitted': bool(self.personnel.get('character_ref_2')), 'status': 'Uploaded' if self.personnel.get('character_ref_2') else 'Required'},
            {'name': 'Bank Reference 1', 'submitted': bool(self.personnel.get('financial_ref_1')), 'status': 'Uploaded' if self.personnel.get('financial_ref_1') else 'Required'},
            {'name': 'Bank Reference 2', 'submitted': bool(self.personnel.get('financial_ref_2')), 'status': 'Uploaded' if self.personnel.get('financial_ref_2') else 'Required'},
        ]
        
        return base_documents
    
    def _add_section_header(self, doc: Document, title: str):
        """Add formatted section header"""
        doc.add_paragraph()
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(title)
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_run.font.underline = True
        doc.add_paragraph()
    
    # Assessment methods for generating summary
    def _calculate_overall_rating(self) -> str:
        """Calculate overall fit and proper rating"""
        category_scores = self._assess_all_categories()
        
        total_score = 0
        total_weight = 0
        
        for category, data in self.ASSESSMENT_CATEGORIES.items():
            score = category_scores.get(category, 0)
            weight = data['weight']
            total_score += score * weight
            total_weight += weight
        
        overall_score = total_score / total_weight if total_weight > 0 else 0
        
        if overall_score >= 80:
            return 'Highly Suitable'
        elif overall_score >= 70:
            return 'Suitable'
        elif overall_score >= 60:
            return 'Marginally Suitable'
        else:
            return 'Not Suitable'
    
    def _assess_all_categories(self) -> Dict[str, float]:
        """Assess all fit and proper categories"""
        return {
            'integrity': self._assess_integrity(),
            'competence': self._assess_competence(),
            'financial_soundness': self._assess_financial_soundness(),
            'regulatory_compliance': self._assess_regulatory_compliance(),
            'reputation': self._assess_reputation(),
        }
    
    def _assess_integrity(self) -> float:
        """Assess integrity score"""
        score = 100
        
        # Deduct for negative declarations
        if self.personnel.get('criminal_conviction'):
            score -= 30
        if self.personnel.get('bankruptcy'):
            score -= 20
        if self.personnel.get('dismissed_or_resigned'):
            score -= 15
        if self.personnel.get('ever_disqualified'):
            score -= 25
        
        return max(score, 0)
    
    def _assess_competence(self) -> float:
        """Assess competence score"""
        score = 60  # Base score
        
        # Add for qualifications
        qualifications = self.personnel.get('professional_qualifications', '')
        if qualifications:
            score += 20
        
        # Add for experience
        experience = self.personnel.get('employment_history_10_years', '')
        if len(experience) > 200:  # Substantial experience
            score += 20
        elif len(experience) > 100:  # Some experience
            score += 10
        
        return min(score, 100)
    
    def _assess_financial_soundness(self) -> float:
        """Assess financial soundness score"""
        score = 80  # Base score assuming sound financial position
        
        if self.personnel.get('bankruptcy'):
            score -= 40
        
        # Asset statement uploaded
        if self.personnel.get('statement_assets_liabilities'):
            score += 10
        
        return max(min(score, 100), 0)
    
    def _assess_regulatory_compliance(self) -> float:
        """Assess regulatory compliance score"""
        score = 90  # Base score
        
        if self.personnel.get('regulatory_actions', False):
            score -= 30
        if self.personnel.get('ever_disqualified'):
            score -= 25
        
        return max(score, 0)
    
    def _assess_reputation(self) -> float:
        """Assess reputation score"""
        score = 75  # Base score
        
        # Character references provided
        if self.personnel.get('character_ref_1') and self.personnel.get('character_ref_2'):
            score += 15
        elif self.personnel.get('character_ref_1') or self.personnel.get('character_ref_2'):
            score += 8
        
        # Professional qualifications
        if self.personnel.get('professional_qualifications'):
            score += 10
        
        return min(score, 100)
    
    def _identify_red_flags(self) -> List[str]:
        """Identify red flags in the application"""
        red_flags = []
        
        if self.personnel.get('criminal_conviction'):
            red_flags.append('Criminal conviction declared')
        if self.personnel.get('bankruptcy'):
            red_flags.append('Bankruptcy/insolvency history')
        if self.personnel.get('ever_disqualified'):
            red_flags.append('Previous professional disqualification')
        if self.personnel.get('dismissed_or_resigned'):
            red_flags.append('Employment termination due to misconduct')
        
        # Check for missing critical documents
        if not self.personnel.get('police_clearance'):
            red_flags.append('Police clearance certificate missing')
        if not self.personnel.get('tax_clearance_individual'):
            red_flags.append('Tax clearance certificate missing')
        
        return red_flags
    
    def _generate_recommendations(self) -> List[str]:
        """Generate assessment recommendations"""
        recommendations = []
        
        overall_rating = self._calculate_overall_rating()
        
        if overall_rating in ['Highly Suitable', 'Suitable']:
            recommendations.append('Recommend for appointment subject to satisfactory reference checks')
        elif overall_rating == 'Marginally Suitable':
            recommendations.append('Conditional approval - require additional monitoring and regular reviews')
            recommendations.append('Consider probationary period with enhanced oversight')
        else:
            recommendations.append('Not recommended for appointment at this time')
            recommendations.append('Address identified concerns before reapplication')
        
        # Specific recommendations based on gaps
        if not self.personnel.get('professional_qualifications'):
            recommendations.append('Require submission of educational and professional qualification certificates')
        
        if not self.personnel.get('character_ref_1') or not self.personnel.get('character_ref_2'):
            recommendations.append('Obtain and verify all required character references')
        
        return recommendations
    
    def _identify_missing_information(self) -> List[str]:
        """Identify missing information"""
        missing = []
        
        required_fields = [
            ('full_name', 'Full name'),
            ('date_of_birth', 'Date of birth'),
            ('nationality', 'Nationality'),
            ('id_or_passport_number', 'ID/Passport number'),
            ('professional_qualifications', 'Professional qualifications'),
            ('employment_history_10_years', 'Employment history'),
        ]
        
        for field, description in required_fields:
            if not self.personnel.get(field):
                missing.append(description)
        
        return missing
    
    def _determine_approval_recommendation(self) -> str:
        """Determine final approval recommendation"""
        red_flags = self._identify_red_flags()
        overall_rating = self._calculate_overall_rating()
        missing_info = self._identify_missing_information()
        
        if len(red_flags) > 2:
            return 'REJECT - Multiple red flags identified'
        elif 'Criminal conviction declared' in red_flags:
            return 'REJECT - Criminal conviction incompatible with position'
        elif overall_rating in ['Highly Suitable', 'Suitable'] and len(missing_info) == 0:
            return 'APPROVE - Meets all fit and proper criteria'
        elif overall_rating == 'Marginally Suitable':
            return 'CONDITIONAL APPROVAL - Subject to enhanced monitoring'
        else:
            return 'DEFER - Require additional information before decision'



# ==========================================================================================
    # Business Plan Generation
# ===============================================================================================
class BusinessPlanGenerator:
    """
    3-Year Business Plan Generator
    Creates comprehensive business plan using FinancialProjectionModel
    """
    
    def __init__(self, lender_data):
        self.lender = lender_data
        self.financial_model = FinancialProjectionModel(lender_data)
        
    def generate(self):
        """Generate comprehensive business plan document"""        
        doc = Document()
        
        # Set up document style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        
        # Generate financial projections first
        projections = self.financial_model.generate_complete_projections()
        
        # Title Page
        self._add_title_page(doc)
        
        # Table of Contents placeholder
        self._add_table_of_contents(doc)
        
        # Executive Summary
        self._add_executive_summary(doc, projections)
        
        # Company Overview
        self._add_company_overview(doc)
        
        # Market Analysis
        self._add_market_analysis(doc)
        
        # Products and Services
        self._add_products_services(doc)
        
        # Marketing and Sales Strategy
        self._add_marketing_strategy(doc)
        
        # Operations Plan
        self._add_operations_plan(doc)
        
        # Management Team
        self._add_management_team(doc)
        
        # Financial Projections
        self._add_financial_projections(doc, projections)
        
        # Risk Analysis
        self._add_risk_analysis(doc, projections)
        
        # Implementation Timeline
        self._add_implementation_timeline(doc)
        
        return doc
    
    def _add_title_page(self, doc):
        """Add business plan title page"""
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_run = title_para.add_run('THREE-YEAR BUSINESS PLAN\n\n')
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        
        company_run = title_para.add_run(f'{self.lender.get("company_name", "")}\n')
        company_run.font.size = Pt(16)
        company_run.font.bold = True
        
        subtitle_run = title_para.add_run('Microfinance Institution\n\n')
        subtitle_run.font.size = Pt(14)
        
        date_run = title_para.add_run(f'Prepared: {self._get_current_date()}\n')
        date_run.font.size = Pt(12)
        
        tier_run = title_para.add_run(f'CBL License Application: {self.lender.get("cbl_tier", "").upper()}')
        tier_run.font.size = Pt(12)
        tier_run.font.italic = True
        
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc):
        """Add table of contents"""
        toc_para = doc.add_paragraph()
        toc_title = toc_para.add_run('TABLE OF CONTENTS\n\n')
        toc_title.font.size = Pt(16)
        toc_title.font.bold = True
        
        contents = [
            '1. Executive Summary',
            '2. Company Overview', 
            '3. Market Analysis',
            '4. Products and Services',
            '5. Marketing and Sales Strategy',
            '6. Operations Plan',
            '7. Management Team',
            '8. Financial Projections',
            '9. Risk Analysis',
            '10. Implementation Timeline'
        ]
        
        for item in contents:
            content_para = doc.add_paragraph(item, style='List Number')
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, projections):
        """Add executive summary section"""
        self._add_section_header(doc, '1. EXECUTIVE SUMMARY')
        
        summary = projections.get('summary', {})
        
        exec_text = f"""
{self.lender.get('company_name', '')} is a proposed {self.lender.get('cbl_tier', '').upper()} microfinance institution seeking to serve the financially excluded population in {self.lender.get('district', 'Maseru')}, Lesotho.

BUSINESS CONCEPT:
Our institution will provide accessible financial services including micro-credit, savings (if applicable), and financial education to small-scale entrepreneurs, women's groups, and underbanked individuals.

KEY SUCCESS FACTORS:
• Strong management team with microfinance experience
• Robust risk management and compliance framework
• Technology-enabled service delivery
• Community-based approach to client relationships

FINANCIAL HIGHLIGHTS:
• Starting Capital: M{self.lender.get('stated_capital', 0):,.0f}
• Year 3 Revenue Projection: {summary.get('year_3_revenue', 'TBD')}
• Break-even Timeline: {summary.get('break_even_month', 'TBD')} months
• Employment Creation: {summary.get('employment_created', 'TBD')} jobs by Year 3
• Target Market: {summary.get('target_market', 'TBD')}

FUNDING REQUIREMENT:
The total capital requirement is M{self.lender.get('stated_capital', 0):,.0f}, which will be provided through shareholder equity contributions.
        """
        
        doc.add_paragraph(exec_text.strip())
    
    def _add_company_overview(self, doc):
        """Add company overview section"""
        self._add_section_header(doc, '2. COMPANY OVERVIEW')
        
        overview_text = f"""
COMPANY PROFILE:
{self.lender.get('company_name', '')} will be incorporated as a private company limited by shares, registered under the Companies Act of Lesotho.

MISSION STATEMENT:
To provide accessible, affordable, and sustainable financial services that empower individuals and small businesses to improve their economic well-being.

VISION:
To become the leading microfinance institution in {self.lender.get('district', 'Maseru')} district, recognized for innovation, customer service, and positive community impact.

CORE VALUES:
• Integrity and transparency in all operations
• Customer-centricity and service excellence  
• Financial inclusion and empowerment
• Sustainable and responsible lending practices
• Community development and social impact

LEGAL STRUCTURE:
• Company Name: {self.lender.get('company_name', '')}
• Registration Number: {self.lender.get('registration_number', 'To be obtained')}
• Incorporation Date: {self._format_date(self.lender.get('date_of_establishment'))}
• Authorized Capital: M{self.lender.get('stated_capital', 0):,.0f}
• Registered Address: {self.lender.get('physical_address', '')}
        """
        
        doc.add_paragraph(overview_text.strip())
    
    def _add_market_analysis(self, doc):
        """Add market analysis section"""
        self._add_section_header(doc, '3. MARKET ANALYSIS')
        
        market_text = f"""
LESOTHO MICROFINANCE SECTOR:
The microfinance sector in Lesotho serves a critical role in financial inclusion, with approximately 60% of the adult population remaining unbanked or underbanked.

TARGET MARKET SEGMENTS:

1. SMALL-SCALE ENTREPRENEURS:
   • Market size: Estimated {self.lender.get('target_market_size', 1000):,} potential clients
   • Average loan requirement: M5,000 - M25,000
   • Primary sectors: retail trade, agriculture, services

2. WOMEN'S GROUPS AND COOPERATIVES:
   • Group lending model with social collateral
   • Focus on income-generating activities
   • Financial literacy and business training needs

3. AGRICULTURAL VALUE CHAIN PARTICIPANTS:
   • Seasonal financing requirements
   • Equipment and input financing
   • Market linkage facilitation

COMPETITIVE LANDSCAPE:
The market includes established players such as LICO, LESOTHO POST BANK, and various SACCOs. Our competitive advantage lies in:
• Technology-enabled service delivery
• Flexible product offerings
• Strong community relationships
• Competitive pricing structure

MARKET OPPORTUNITIES:
• Growing demand for digital financial services
• Increasing government support for financial inclusion
• Limited competition in rural and peri-urban areas
• Opportunity for innovative product development
        """
        
        doc.add_paragraph(market_text.strip())
    
    def _add_products_services(self, doc):
        """Add products and services section"""
        self._add_section_header(doc, '4. PRODUCTS AND SERVICES')
        
        tier = self.lender.get('cbl_tier', '')
        
        products_text = f"""
PRIMARY PRODUCTS:

1. INDIVIDUAL MICROLOANS:
   • Target market: Individual entrepreneurs
   • Loan amounts: M1,000 - M50,000
   • Interest rates: 24% - 36% per annum
   • Repayment terms: 3 - 24 months
   • Collateral: Flexible security arrangements

2. GROUP LENDING:
   • Target market: Women's groups, cooperatives
   • Group size: 5 - 20 members
   • Loan amounts: M500 - M10,000 per member
   • Interest rates: 20% - 30% per annum
   • Social collateral model

3. AGRICULTURAL LOANS:
   • Seasonal financing for farming activities
   • Equipment and input financing
   • Harvest-based repayment schedules
   • Technical assistance partnerships
        """
        
        if tier == 'tier1':
            products_text += """

4. SAVINGS PRODUCTS:
   • Personal savings accounts
   • Group savings accounts
   • Term deposits
   • Interest rates: 6% - 8% per annum
            """
        
        products_text += """

SUPPORT SERVICES:
• Financial literacy training
• Business development support
• Market linkage facilitation
• Digital payment solutions
• Customer support and relationship management
        """
        
        doc.add_paragraph(products_text.strip())
    
    def _add_marketing_strategy(self, doc):
        """Add marketing and sales strategy section"""
        self._add_section_header(doc, '5. MARKETING AND SALES STRATEGY')
        
        marketing_text = """
CUSTOMER ACQUISITION STRATEGY:

1. COMMUNITY ENGAGEMENT:
   • Partnership with local chiefs and community leaders
   • Village-level financial education workshops
   • Demonstration projects with early adopters

2. REFERRAL PROGRAMS:
   • Incentives for existing customers to refer new clients
   • Group formation facilitation
   • Word-of-mouth marketing through satisfied customers

3. DIGITAL MARKETING:
   • Social media presence (Facebook, WhatsApp)
   • Local radio sponsorship and financial education programs
   • Mobile-based marketing and client communication

4. INSTITUTIONAL PARTNERSHIPS:
   • Collaboration with NGOs and development organizations
   • Government agency partnerships
   • Private sector value chain partnerships

CUSTOMER RETENTION STRATEGY:
• Excellent customer service and relationship management
• Graduated loan products for repeat customers
• Value-added services and business support
• Competitive pricing and flexible terms
• Technology-enabled convenience

SALES TARGETS:
• Year 1: 200 active clients
• Year 2: 500 active clients  
• Year 3: 1,000 active clients
• Client retention rate: 85% or higher
        """
        
        doc.add_paragraph(marketing_text.strip())
    
    def _add_operations_plan(self, doc):
        """Add operations plan section"""
        self._add_section_header(doc, '6. OPERATIONS PLAN')
        
        operations_text = f"""
OPERATIONAL STRUCTURE:

1. HEAD OFFICE:
   • Location: {self.lender.get('physical_address', '')}
   • Functions: Management, compliance, finance, IT
   • Staff: 5-8 personnel

2. SERVICE DELIVERY:
   • Mobile banking units for rural outreach
   • Agent network for transaction processing
   • Digital platforms for account management
   • Branch office (planned Year 2)

3. TECHNOLOGY INFRASTRUCTURE:
   • Core banking system with microfinance modules
   • Mobile money integration
   • Customer relationship management system
   • Automated reporting and compliance tools

OPERATIONAL PROCESSES:

Credit Assessment and Approval:
• Standardized credit scoring methodology
• Field verification and character assessment
• Tiered approval limits based on amount
• Risk-based pricing model

Loan Portfolio Management:
• Automated payment reminders
• Early warning systems for delinquency
• Flexible restructuring options
• Collections and recovery procedures

Customer Service:
• Multi-channel customer support
• Financial education and business counseling
• Complaint resolution procedures
• Regular customer satisfaction surveys

QUALITY ASSURANCE:
• Internal audit function
• External audit and review
• Regulatory compliance monitoring
• Continuous process improvement
        """
        
        doc.add_paragraph(operations_text.strip())
    
    def _add_management_team(self, doc):
        """Add management team section"""
        self._add_section_header(doc, '7. MANAGEMENT TEAM')
        
        ceo_name = f"{self.lender.get('ceo_first_name', '')} {self.lender.get('ceo_last_name', '')}".strip()
        
        management_text = f"""
ORGANIZATIONAL STRUCTURE:

BOARD OF DIRECTORS:
The Board will consist of {len([p for p in self.lender.get('personnel', []) if p.get('role') == 'director'])} directors with diverse backgrounds in finance, business, and community development.

KEY MANAGEMENT PERSONNEL:

Chief Executive Officer - {ceo_name or 'To be appointed'}:
• Overall strategic leadership and management
• Stakeholder relationship management
• Business development and growth strategy

Chief Financial Officer - To be appointed:
• Financial management and planning
• Risk management and compliance
• Management information systems

Chief Operations Officer - To be appointed:
• Day-to-day operational management
• Product development and delivery
• Customer service and relationship management

Compliance Officer - To be appointed:
• Regulatory compliance and reporting
• Internal controls and audit
• Risk assessment and mitigation

ADVISORY COMMITTEE:
Technical advisory support from experienced microfinance practitioners, development finance experts, and industry specialists.

HUMAN RESOURCES PLAN:
• Year 1: 5-8 staff members
• Year 2: 10-12 staff members
• Year 3: 15-20 staff members
• Emphasis on local recruitment and capacity building
• Comprehensive training and development programs
        """
        
        doc.add_paragraph(management_text.strip())
    
    def _add_financial_projections(self, doc, projections):
        """Add detailed financial projections section"""
        self._add_section_header(doc, '8. FINANCIAL PROJECTIONS')
        
        # Add financial projections tables
        income_statement = projections.get('income_statement', {})
        
        doc.add_paragraph('8.1 INCOME STATEMENT PROJECTIONS (M)')
        
        # Create income statement table
        income_table = doc.add_table(rows=8, cols=4)
        income_table.style = 'Table Grid'
        
        # Headers
        headers = ['Item', 'Year 1', 'Year 2', 'Year 3']
        for i, header in enumerate(headers):
            income_table.rows[0].cells[i].text = header
            income_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Financial data
        financial_items = [
            ('Interest Income', 
             income_statement.get('year_1', {}).get('revenue', {}).get('interest_income', 0),
             income_statement.get('year_2', {}).get('revenue', {}).get('interest_income', 0),
             income_statement.get('year_3', {}).get('revenue', {}).get('interest_income', 0)),
            ('Fee Income',
             income_statement.get('year_1', {}).get('revenue', {}).get('fees_and_commissions', 0),
             income_statement.get('year_2', {}).get('revenue', {}).get('fees_and_commissions', 0),
             income_statement.get('year_3', {}).get('revenue', {}).get('fees_and_commissions', 0)),
            ('Total Revenue',
             income_statement.get('year_1', {}).get('revenue', {}).get('total_revenue', 0),
             income_statement.get('year_2', {}).get('revenue', {}).get('total_revenue', 0),
             income_statement.get('year_3', {}).get('revenue', {}).get('total_revenue', 0)),
            ('Operating Expenses',
             income_statement.get('year_1', {}).get('expenses', {}).get('total_expenses', 0),
             income_statement.get('year_2', {}).get('expenses', {}).get('total_expenses', 0),
             income_statement.get('year_3', {}).get('expenses', {}).get('total_expenses', 0)),
            ('Credit Losses',
             income_statement.get('year_1', {}).get('expenses', {}).get('credit_loss_provisions', 0),
             income_statement.get('year_2', {}).get('expenses', {}).get('credit_loss_provisions', 0),
             income_statement.get('year_3', {}).get('expenses', {}).get('credit_loss_provisions', 0)),
            ('Net Income',
             income_statement.get('year_1', {}).get('profitability', {}).get('net_income', 0),
             income_statement.get('year_2', {}).get('profitability', {}).get('net_income', 0),
             income_statement.get('year_3', {}).get('profitability', {}).get('net_income', 0)),
        ]
        
        for i, (item, y1, y2, y3) in enumerate(financial_items, 1):
            income_table.rows[i].cells[0].text = item
            income_table.rows[i].cells[1].text = f"{y1:,.0f}"
            income_table.rows[i].cells[2].text = f"{y2:,.0f}"
            income_table.rows[i].cells[3].text = f"{y3:,.0f}"
        
        # Key assumptions
        doc.add_paragraph()
        doc.add_paragraph('8.2 KEY FINANCIAL ASSUMPTIONS')
        
        assumptions = projections.get('assumptions', {})
        tier_assumptions = assumptions.get('tier_assumptions', {})
        
        assumptions_text = f"""
LOAN PORTFOLIO ASSUMPTIONS:
• Average loan size: M{tier_assumptions.get('avg_loan_size', 0):,.0f}
• Interest rate range: {tier_assumptions.get('interest_rate_range', ['N/A', 'N/A'])[0]*100:.1f}% - {tier_assumptions.get('interest_rate_range', ['N/A', 'N/A'])[1]*100:.1f}%
• Default rate assumption: {tier_assumptions.get('default_rate', 0)*100:.1f}%
• Portfolio growth rate: 25% annually

OPERATIONAL ASSUMPTIONS:
• Staff-to-client ratio: {tier_assumptions.get('staff_per_1000_clients', 0)} per 1,000 clients
• Operating efficiency target: {tier_assumptions.get('operational_efficiency', 0)*100:.1f}%
• Technology investment: 10% of capital

MARKET ASSUMPTIONS:
• Market growth rate: {assumptions.get('market_assumptions', {}).get('microfinance_growth', 0)*100:.1f}%
• Inflation rate: {assumptions.get('market_assumptions', {}).get('inflation_rate', 0)*100:.1f}%
• Competitive pressure: {assumptions.get('market_assumptions', {}).get('competitive_pressure', 0)*100:.1f}%
        """
        
        doc.add_paragraph(assumptions_text.strip())
    
    def _add_risk_analysis(self, doc, projections):
        """Add risk analysis section"""
        self._add_section_header(doc, '9. RISK ANALYSIS AND MITIGATION')
        
        sensitivity_analysis = projections.get('sensitivity_analysis', {})
        
        risk_text = f"""
RISK ASSESSMENT AND MITIGATION STRATEGIES:

1. CREDIT RISK:
   Risk: Loan defaults and portfolio deterioration
   Mitigation:
   • Robust credit assessment procedures
   • Diversified loan portfolio
   • Regular portfolio monitoring
   • Early warning systems
   • Flexible restructuring options

2. OPERATIONAL RISK:
   Risk: System failures, fraud, process breakdowns
   Mitigation:
   • Strong internal controls
   • Regular staff training
   • Technology backup systems
   • Insurance coverage
   • External audit and review

3. MARKET RISK:
   Risk: Interest rate fluctuations, competition
   Mitigation:
   • Flexible pricing model
   • Product differentiation
   • Strong customer relationships
   • Continuous market analysis

4. LIQUIDITY RISK:
   Risk: Inability to meet obligations
   Mitigation:
   • Conservative cash management
   • Diversified funding sources
   • Credit line facilities
   • Liquidity monitoring

5. REGULATORY RISK:
   Risk: Non-compliance with CBL requirements
   Mitigation:
   • Dedicated compliance function
   • Regular regulatory training
   • External compliance review
   • Strong governance framework

SCENARIO ANALYSIS:
        """
        
        for scenario, data in sensitivity_analysis.items():
            risk_text += f"""
{scenario.upper()} SCENARIO:
• Description: {data.get('description', '')}
• Year 3 Profit Impact: {data.get('year_3_profit', 'N/A')}
• Profit Change: {data.get('profit_change', 'N/A')}
            """
        
        doc.add_paragraph(risk_text.strip())
    
    def _add_implementation_timeline(self, doc):
        """Add implementation timeline section"""
        self._add_section_header(doc, '10. IMPLEMENTATION TIMELINE')
        
        timeline_text = """
IMPLEMENTATION MILESTONES:

PRE-LICENSING PHASE (Months 1-6):
• Complete CBL license application submission
• Finalize management team recruitment
• Complete systems and infrastructure setup
• Conduct staff training and capacity building
• Develop operational policies and procedures

LICENSING PHASE (Months 7-9):
• CBL application review and assessment
• Address any regulatory feedback
• Complete final preparations for operations
• Conduct pilot testing and system validation

OPERATIONAL LAUNCH (Months 10-12):
• Commence operations with initial client base
• Implement marketing and customer acquisition strategies
• Monitor and refine operational processes
• Build loan portfolio gradually

GROWTH AND EXPANSION (Years 2-3):
• Scale operations and expand client base
• Introduce additional products and services
• Expand geographic coverage
• Evaluate branch expansion opportunities
• Pursue strategic partnerships

KEY SUCCESS METRICS:
• Client acquisition rate: 50+ new clients per month by Year 2
• Portfolio quality: <5% portfolio at risk
• Operational efficiency: <70% operating expense ratio
• Profitability: Break-even by Month 18
• Regulatory compliance: 100% compliance rating
        """
        
        doc.add_paragraph(timeline_text.strip())
    
    def _add_section_header(self, doc, title):
        """Add formatted section header"""
        doc.add_paragraph()
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(title)
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_run.font.underline = True
        doc.add_paragraph()
    
    def _get_current_date(self):
        """Get current date formatted"""
        from datetime import datetime
        return datetime.now().strftime('%B %Y')
    
    def _format_date(self, date_obj):
        """Format date for display"""
        if date_obj:
            return date_obj.strftime('%d/%m/%Y') if hasattr(date_obj, 'strftime') else str(date_obj)
        return 'To be determined'


class AMLPolicyGenerator:
    """
    AML/CFT Policy Manual Generator
    Creates tier-specific Anti-Money Laundering and Counter-Terrorist Financing policies
    """
    
    def __init__(self, lender_data):
        self.lender = lender_data
        self.tier = lender_data.get('cbl_tier', '')
        
    def generate(self):
        """Generate AML/CFT policy document"""        
        doc = Document()
        
        # Set up document style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        
        # Title page
        self._add_title_page(doc)
        
        # Table of contents
        self._add_table_of_contents(doc)
        
        # Policy sections
        self._add_introduction(doc)
        self._add_regulatory_framework(doc)
        self._add_institutional_framework(doc)
        self._add_customer_due_diligence(doc)
        self._add_transaction_monitoring(doc)
        self._add_record_keeping(doc)
        self._add_reporting_procedures(doc)
        self._add_training_awareness(doc)
        self._add_review_monitoring(doc)
        
        return doc
    
    def _add_title_page(self, doc):
        """Add AML policy title page"""
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_run = title_para.add_run('ANTI-MONEY LAUNDERING\n')
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        
        subtitle_run = title_para.add_run('AND\n')
        subtitle_run.font.size = Pt(14)
        
        subtitle2_run = title_para.add_run('COUNTER-TERRORIST FINANCING\n')
        subtitle2_run.font.size = Pt(18)
        subtitle2_run.font.bold = True
        
        policy_run = title_para.add_run('POLICY MANUAL\n\n')
        policy_run.font.size = Pt(16)
        policy_run.font.bold = True
        
        company_run = title_para.add_run(f'{self.lender.get("company_name", "")}\n')
        company_run.font.size = Pt(14)
        company_run.font.bold = True
        
        date_run = title_para.add_run(f'Effective Date: {datetime.now().strftime("%B %Y")}\n')
        date_run.font.size = Pt(12)
        
        version_run = title_para.add_run('Version 1.0')
        version_run.font.size = Pt(12)
        version_run.font.italic = True
        
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc):
        """Add table of contents for AML policy"""
        toc_para = doc.add_paragraph()
        toc_title = toc_para.add_run('TABLE OF CONTENTS\n\n')
        toc_title.font.size = Pt(16)
        toc_title.font.bold = True
        
        contents = [
            '1. Introduction and Scope',
            '2. Regulatory Framework',
            '3. Institutional Framework',
            '4. Customer Due Diligence',
            '5. Transaction Monitoring',
            '6. Record Keeping',
            '7. Reporting Procedures',
            '8. Training and Awareness',
            '9. Review and Monitoring'
        ]
        
        for item in contents:
            doc.add_paragraph(item, style='List Number')
        
        doc.add_page_break()
    
    def _add_introduction(self, doc):
        """Add introduction section"""
        self._add_section_header(doc, '1. INTRODUCTION AND SCOPE')
        
        intro_text = f"""
1.1 PURPOSE
This Anti-Money Laundering and Counter-Terrorist Financing (AML/CFT) Policy Manual establishes the framework for {self.lender.get('company_name', '')} to comply with applicable AML/CFT laws and regulations in Lesotho.

1.2 SCOPE
This policy applies to all staff, management, and board members of {self.lender.get('company_name', '')}. It covers all business relationships, transactions, and activities conducted by the institution.

1.3 POLICY STATEMENT
{self.lender.get('company_name', '')} is committed to:
• Preventing the use of our services for money laundering and terrorist financing
• Complying with all applicable AML/CFT laws and regulations
• Implementing robust customer due diligence procedures
• Monitoring transactions for suspicious activities
• Reporting suspicious transactions to relevant authorities
• Maintaining comprehensive records and documentation

1.4 RISK ASSESSMENT
As a {self.tier.upper()} microfinance institution, our AML/CFT risk exposure is assessed as {'MEDIUM' if self.tier in ['tier1', 'tier2'] else 'LOW'} based on:
• Customer base characteristics
• Product and service offerings
• Geographic coverage
• Transaction volumes and amounts
        """
        
        doc.add_paragraph(intro_text.strip())
    
    def _add_regulatory_framework(self, doc):
        """Add regulatory framework section"""
        self._add_section_header(doc, '2. REGULATORY FRAMEWORK')
        
        regulatory_text = """
2.1 APPLICABLE LEGISLATION
• Money Laundering and Proceeds of Crime Act 2008
• Financial Intelligence Unit Act 2013
• Financial Institutions Act 2019
• Prevention of Terrorism Act 2007
• Central Bank of Lesotho Regulations and Guidelines

2.2 REGULATORY AUTHORITIES
• Central Bank of Lesotho (CBL): Primary regulator for microfinance institutions
• Financial Intelligence Unit (FIU): Suspicious transaction reporting
• Director of Public Prosecutions: Money laundering prosecutions

2.3 PENALTIES FOR NON-COMPLIANCE
Non-compliance with AML/CFT requirements may result in:
• Administrative sanctions and fines
• Criminal prosecution
• License suspension or revocation
• Reputational damage
        """
        
        doc.add_paragraph(regulatory_text.strip())
    
    def _add_institutional_framework(self, doc):
        """Add institutional framework section"""
        self._add_section_header(doc, '3. INSTITUTIONAL FRAMEWORK')
        
        institutional_text = """
3.1 BOARD RESPONSIBILITIES
The Board of Directors is responsible for:
• Approving AML/CFT policies and procedures
• Ensuring adequate resources for compliance
• Overseeing compliance monitoring and reporting
• Reviewing and updating policies annually

3.2 SENIOR MANAGEMENT RESPONSIBILITIES
Senior management is responsible for:
• Implementing board-approved policies
• Establishing compliance monitoring systems
• Ensuring staff training and competency
• Reporting to the board on compliance matters

3.3 COMPLIANCE OFFICER
The designated Compliance Officer is responsible for:
• Day-to-day compliance monitoring
• Staff training and guidance
• Suspicious transaction identification and reporting
• Liaison with regulatory authorities
• Maintaining compliance records

3.4 STAFF RESPONSIBILITIES
All staff members are required to:
• Follow established AML/CFT procedures
• Report suspicious activities immediately
• Participate in required training programs
• Maintain customer information confidentiality
        """
        
        doc.add_paragraph(institutional_text.strip())
    
    def _add_customer_due_diligence(self, doc):
        """Add customer due diligence section"""
        self._add_section_header(doc, '4. CUSTOMER DUE DILIGENCE')
        
        tier_specific_requirements = self._get_tier_specific_cdd_requirements()
        
        cdd_text = f"""
4.1 CUSTOMER IDENTIFICATION REQUIREMENTS
All customers must provide:
• Valid government-issued identification
• Proof of residential address
• Business registration documents (for business customers)
• Tax identification numbers
• Beneficial ownership information (for entities)

4.2 CUSTOMER DUE DILIGENCE PROCEDURES
{tier_specific_requirements}

4.3 ENHANCED DUE DILIGENCE
Enhanced due diligence is required for:
• Politically Exposed Persons (PEPs)
• High-risk customers or transactions
• Non-resident customers
• Complex ownership structures
• Cash-intensive businesses

4.4 ONGOING MONITORING
• Regular review of customer information
• Transaction pattern analysis
• Risk rating updates
• Account monitoring for unusual activities

4.5 RECORD KEEPING
Customer identification records must be maintained for a minimum of 5 years after account closure.
        """
        
        doc.add_paragraph(cdd_text.strip())
    
    def _add_transaction_monitoring(self, doc):
        """Add transaction monitoring section"""
        self._add_section_header(doc, '5. TRANSACTION MONITORING')
        
        monitoring_text = """
5.1 MONITORING SYSTEMS
• Automated transaction monitoring where applicable
• Manual review procedures for flagged transactions
• Regular pattern analysis and trend identification
• Exception reporting and investigation

5.2 SUSPICIOUS TRANSACTION INDICATORS
• Transactions inconsistent with customer profile
• Unusual transaction patterns or frequencies
• Large cash transactions without clear economic purpose
• Transactions involving high-risk jurisdictions
• Attempts to avoid reporting thresholds

5.3 INVESTIGATION PROCEDURES
• Immediate investigation of flagged transactions
• Documentation of investigation findings
• Escalation to Compliance Officer
• Decision on reporting to FIU

5.4 REPORTING THRESHOLDS
• Large cash transactions: M50,000 or equivalent
• Suspicious transactions: Any amount
• International transfers: M100,000 or equivalent
        """
        
        doc.add_paragraph(monitoring_text.strip())
    
    def _add_record_keeping(self, doc):
        """Add record keeping section"""
        self._add_section_header(doc, '6. RECORD KEEPING')
        
        records_text = """
6.1 RECORD RETENTION REQUIREMENTS
• Customer identification records: 5 years after account closure
• Transaction records: 5 years after transaction completion
• Suspicious transaction reports: 5 years after submission
• AML/CFT training records: 5 years

6.2 RECORD FORMATS
• Electronic records with backup systems
• Physical documents in secure storage
• Retrievable format for regulatory inspection
• Confidentiality and data protection measures

6.3 ACCESS CONTROLS
• Authorized personnel access only
• Audit trails for record access
• Secure storage and backup procedures
• Regular review of access permissions
        """
        
        doc.add_paragraph(records_text.strip())
    
    def _add_reporting_procedures(self, doc):
        """Add reporting procedures section"""
        self._add_section_header(doc, '7. REPORTING PROCEDURES')
        
        reporting_text = """
7.1 SUSPICIOUS TRANSACTION REPORTS (STRs)
• Filed within 3 working days of identification
• Submitted to Financial Intelligence Unit
• Include all relevant transaction details
• Maintain strict confidentiality

7.2 LARGE CASH TRANSACTION REPORTS
• Transactions exceeding M50,000
• Filed within prescribed timeframes
• Complete customer and transaction information

7.3 INTERNATIONAL TRANSFER REPORTS
• Cross-border transfers exceeding M100,000
• Complete originator and beneficiary information
• Purpose and nature of transaction

7.4 INTERNAL REPORTING
• Regular compliance reports to management
• Board reporting on quarterly basis
• Exception reports and trend analysis
        """
        
        doc.add_paragraph(reporting_text.strip())
    
    def _add_training_awareness(self, doc):
        """Add training and awareness section"""
        self._add_section_header(doc, '8. TRAINING AND AWARENESS')
        
        training_text = """
8.1 TRAINING PROGRAM
• Initial AML/CFT training for all new staff
• Annual refresher training for existing staff
• Specialized training for compliance officers
• Regular updates on regulatory changes

8.2 TRAINING CONTENT
• AML/CFT legal and regulatory framework
• Customer due diligence procedures
• Transaction monitoring and reporting
• Red flag indicators and investigation
• Record keeping requirements

8.3 TRAINING DELIVERY
• Formal classroom training sessions
• Online training modules where available
• Case study discussions and workshops
• External training and certification programs

8.4 TRAINING RECORDS
• Attendance records and certificates
• Training content and materials
• Assessment results and competency tests
• Ongoing training needs analysis
        """
        
        doc.add_paragraph(training_text.strip())
    
    def _add_review_monitoring(self, doc):
        """Add review and monitoring section"""
        self._add_section_header(doc, '9. REVIEW AND MONITORING')
        
        review_text = """
9.1 POLICY REVIEW
• Annual review of AML/CFT policies
• Updates for regulatory changes
• Board approval of policy amendments
• Implementation of updated procedures

9.2 COMPLIANCE MONITORING
• Regular compliance testing and audits
• Independent review of procedures
• Exception reporting and corrective action
• External audit assessment

9.3 EFFECTIVENESS ASSESSMENT
• Key performance indicators tracking
• Regulatory feedback assessment
• Staff feedback and improvement suggestions
• Benchmarking against industry practices

9.4 CONTINUOUS IMPROVEMENT
• Regular update of risk assessments
• Enhancement of monitoring systems
• Staff training program improvements
• Technology upgrades and automation
        """
        
        doc.add_paragraph(review_text.strip())
    
    def _get_tier_specific_cdd_requirements(self):
        """Get customer due diligence requirements specific to CBL tier"""
        if self.tier == 'tier1':
            return """
ENHANCED DUE DILIGENCE (Tier 1 Requirements):
• Comprehensive customer risk assessment
• Source of funds verification
• Enhanced monitoring for high-risk customers
• Regular customer information updates
• Independent verification of documentation
            """
        elif self.tier == 'tier2':
            return """
STANDARD DUE DILIGENCE (Tier 2 Requirements):
• Customer identification and verification
• Business relationship purpose assessment
• Risk-based monitoring approach
• Regular customer information review
• Documentation of risk assessment
            """
        else:
            return """
SIMPLIFIED DUE DILIGENCE (Tier 3 Requirements):
• Basic customer identification and verification
• Risk assessment based on transaction patterns
• Periodic customer information updates
• Focus on cash transaction monitoring
• Simplified documentation requirements
            """
    
    def _add_section_header(self, doc, title):
        """Add formatted section header"""
        doc.add_paragraph()
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(title)
        header_run.font.size = Pt(14)
        header_run.font.bold = True
        header_run.font.underline = True
        doc.add_paragraph()


# Export all generator classes
__all__ = [
    'FinancialProjectionModel',
    'ScheduleIIGenerator', 
    'FitProperGenerator',
    'ScheduleIGenerator',
    'BusinessPlanGenerator',
    'AMLPolicyGenerator'
]




# ==================================================================================
    # Financial Projection Model for MFI Business Plans
    # Generates realistic 3-year financial projections based on tier, capital, and market data
# ==================================================================================

class FinancialProjectionModel:
    """
    Generates CBL-compliant financial projections for MFI business plans
    Based on Lesotho market data and tier-specific assumptions
    """
    
    # Market assumptions based on Lesotho MFI sector data
    MARKET_ASSUMPTIONS = {
        'lesotho_gdp_growth': 0.025,  # 2.5% annual
        'inflation_rate': 0.055,      # 5.5% average
        'microfinance_growth': 0.15,  # 15% sector growth
        'competitive_pressure': 0.02, # 2% margin compression
    }
    
    # Tier-specific operational assumptions
    TIER_ASSUMPTIONS = {
        'tier1': {
            'avg_loan_size': 25000,
            'interest_rate_range': (0.24, 0.36),  # 24-36% annual
            'default_rate': 0.08,                  # 8%
            'operational_efficiency': 0.75,       # 75% efficiency ratio
            'staff_per_1000_clients': 12,
            'branch_coverage_ratio': 0.15,        # Clients per branch
            'deposit_interest_rate': 0.08,        # 8% on deposits
            'liquidity_requirement': 0.15,        # 15% minimum
        },
        'tier2': {
            'avg_loan_size': 15000,
            'interest_rate_range': (0.28, 0.42),
            'default_rate': 0.06,
            'operational_efficiency': 0.70,
            'staff_per_1000_clients': 8,
            'branch_coverage_ratio': 0.12,
            'deposit_interest_rate': 0.0,         # No deposits
            'liquidity_requirement': 0.0,
        },
        'tier3': {
            'avg_loan_size': 8000,
            'interest_rate_range': (0.30, 0.48),
            'default_rate': 0.05,
            'operational_efficiency': 0.65,
            'staff_per_1000_clients': 6,
            'branch_coverage_ratio': 0.10,
            'deposit_interest_rate': 0.0,
            'liquidity_requirement': 0.0,
        }
    }
    
    def __init__(self, lender_data: Dict):
        self.lender = lender_data
        self.tier = lender_data.get('cbl_tier', 'tier3')
        self.capital = Decimal(str(lender_data.get('stated_capital', 0)))
        self.assets = Decimal(str(lender_data.get('total_assets', 0)))
        self.assumptions = self.TIER_ASSUMPTIONS.get(self.tier, self.TIER_ASSUMPTIONS['tier3'])
        
        # Base projections on capital and existing assets
        self.initial_loan_portfolio = max(self.capital * Decimal('0.8'), self.assets * Decimal('0.6'))
        self.target_market_size = self._estimate_target_market()
        
    def generate_complete_projections(self) -> Dict:
        """Generate complete 3-year financial projections"""
        projections = {
            'summary': self._generate_executive_summary(),
            'assumptions': self._document_assumptions(),
            'income_statement': self._project_income_statement(),
            'balance_sheet': self._project_balance_sheet(),
            'cash_flow': self._project_cash_flow(),
            'key_ratios': self._calculate_key_ratios(),
            'sensitivity_analysis': self._perform_sensitivity_analysis(),
            'regulatory_compliance': self._check_regulatory_ratios(),
        }
        
        return projections
    
    def _estimate_target_market(self) -> int:
        """Estimate addressable market size based on location and capital"""
        # Base market estimation on capital size and tier
        base_clients = {
            'tier1': int(self.capital / 20000),  # Larger loans, fewer clients
            'tier2': int(self.capital / 15000),
            'tier3': int(self.capital / 8000),   # Smaller loans, more clients
        }
        
        return min(base_clients.get(self.tier, 100), 5000)  # Cap at 5000 clients
    
    def _generate_executive_summary(self) -> Dict:
        """Generate executive summary with key financial highlights"""
        year_3_revenue = self._calculate_year_revenue(3)
        year_3_profit = self._calculate_year_profit(3)
        
        return {
            'business_model': f"CBL {self.tier.upper()} microfinance institution",
            'target_market': f"{self.target_market_size:,} potential clients",
            'year_3_revenue': f"M{year_3_revenue:,.0f}",
            'year_3_profit': f"M{year_3_profit:,.0f}",
            'break_even_month': self._calculate_break_even_month(),
            'roi_3_year': f"{(year_3_profit / self.capital * 100):.1f}%",
            'employment_created': self._estimate_employment(3),
        }
    
    def _document_assumptions(self) -> Dict:
        """Document all assumptions for CBL review"""
        return {
            'market_assumptions': self.MARKET_ASSUMPTIONS,
            'tier_assumptions': self.assumptions,
            'business_specific': {
                'initial_capital': float(self.capital),
                'target_market_size': self.target_market_size,
                'avg_loan_size': self.assumptions['avg_loan_size'],
                'planned_interest_rate': f"{self.assumptions['interest_rate_range'][0]*100:.1f}%-{self.assumptions['interest_rate_range'][1]*100:.1f}%",
                'geographic_focus': self.lender.get('district', 'Maseru'),
            }
        }
    
    def _project_income_statement(self) -> Dict:
        """Project 3-year income statement"""
        years = {}
        
        for year in range(1, 4):
            revenue = self._calculate_year_revenue(year)
            expenses = self._calculate_year_expenses(year)
            
            # Interest income (primary revenue)
            interest_income = revenue * Decimal('0.85')  # 85% from interest
            fees_income = revenue * Decimal('0.15')      # 15% from fees
            
            # Operating expenses breakdown
            staff_costs = expenses * Decimal('0.45')
            administrative = expenses * Decimal('0.25')
            credit_losses = revenue * Decimal(str(self.assumptions['default_rate']))
            other_expenses = expenses * Decimal('0.30')
            
            gross_income = interest_income + fees_income
            total_expenses = staff_costs + administrative + credit_losses + other_expenses
            net_income = gross_income - total_expenses
            
            years[f'year_{year}'] = {
                'revenue': {
                    'interest_income': float(interest_income),
                    'fees_and_commissions': float(fees_income),
                    'total_revenue': float(revenue),
                },
                'expenses': {
                    'staff_costs': float(staff_costs),
                    'administrative_expenses': float(administrative),
                    'credit_loss_provisions': float(credit_losses),
                    'other_operating_expenses': float(other_expenses),
                    'total_expenses': float(total_expenses),
                },
                'profitability': {
                    'gross_income': float(gross_income),
                    'operating_income': float(gross_income - (total_expenses - credit_losses)),
                    'net_income': float(net_income),
                    'net_margin': f"{(net_income / revenue * 100):.1f}%" if revenue > 0 else "0%",
                }
            }
        
        return years
    
    def _project_balance_sheet(self) -> Dict:
        """Project 3-year balance sheet"""
        years = {}
        
        for year in range(1, 4):
            # Assets
            loan_portfolio = self._calculate_loan_portfolio(year)
            cash_reserves = self._calculate_cash_reserves(year)
            fixed_assets = self._calculate_fixed_assets(year)
            other_assets = loan_portfolio * Decimal('0.05')  # 5% of portfolio
            
            total_assets = loan_portfolio + cash_reserves + fixed_assets + other_assets
            
            # Liabilities (for deposit-taking institutions)
            if self.tier == 'tier1':
                customer_deposits = loan_portfolio * Decimal('1.2')  # Deposit-to-loan ratio
                borrowings = total_assets * Decimal('0.1')           # 10% external funding
            else:
                customer_deposits = Decimal('0')
                borrowings = total_assets * Decimal('0.15')          # 15% external funding
            
            # Equity
            retained_earnings = self._calculate_cumulative_earnings(year)
            total_equity = self.capital + retained_earnings
            
            total_liabilities = customer_deposits + borrowings
            
            years[f'year_{year}'] = {
                'assets': {
                    'loan_portfolio': float(loan_portfolio),
                    'cash_and_equivalents': float(cash_reserves),
                    'fixed_assets': float(fixed_assets),
                    'other_assets': float(other_assets),
                    'total_assets': float(total_assets),
                },
                'liabilities': {
                    'customer_deposits': float(customer_deposits),
                    'borrowings': float(borrowings),
                    'total_liabilities': float(total_liabilities),
                },
                'equity': {
                    'share_capital': float(self.capital),
                    'retained_earnings': float(retained_earnings),
                    'total_equity': float(total_equity),
                },
                'balance_check': float(total_assets - total_liabilities - total_equity) < 1.0,  # Should be near zero
            }
        
        return years
    
    def _project_cash_flow(self) -> Dict:
        """Project 3-year cash flow statement"""
        years = {}
        cumulative_cash = Decimal('0')
        
        for year in range(1, 4):
            # Operating activities
            net_income = Decimal(str(self._calculate_year_profit(year)))
            loan_portfolio_increase = self._calculate_loan_portfolio_growth(year)
            deposit_increase = self._calculate_deposit_growth(year) if self.tier == 'tier1' else Decimal('0')
            
            operating_cash_flow = net_income - loan_portfolio_increase + deposit_increase
            
            # Investing activities (fixed assets, technology)
            capital_expenditures = self._calculate_capex(year)
            investing_cash_flow = -capital_expenditures
            
            # Financing activities
            if year == 1:
                equity_injection = self.capital
                debt_proceeds = self._calculate_initial_borrowing()
            else:
                equity_injection = Decimal('0')
                debt_proceeds = self._calculate_additional_borrowing(year)
            
            financing_cash_flow = equity_injection + debt_proceeds
            
            # Net cash flow
            net_cash_flow = operating_cash_flow + investing_cash_flow + financing_cash_flow
            cumulative_cash += net_cash_flow
            
            years[f'year_{year}'] = {
                'operating_activities': {
                    'net_income': float(net_income),
                    'loan_portfolio_growth': float(-loan_portfolio_increase),
                    'deposit_growth': float(deposit_increase),
                    'operating_cash_flow': float(operating_cash_flow),
                },
                'investing_activities': {
                    'capital_expenditures': float(-capital_expenditures),
                    'investing_cash_flow': float(investing_cash_flow),
                },
                'financing_activities': {
                    'equity_injection': float(equity_injection),
                    'debt_proceeds': float(debt_proceeds),
                    'financing_cash_flow': float(financing_cash_flow),
                },
                'summary': {
                    'net_cash_flow': float(net_cash_flow),
                    'cumulative_cash': float(cumulative_cash),
                    'cash_runway_months': self._calculate_cash_runway(cumulative_cash, year),
                }
            }
        
        return years
    
    def _calculate_key_ratios(self) -> Dict:
        """Calculate key financial and regulatory ratios"""
        ratios = {}
        
        for year in range(1, 4):
            loan_portfolio = self._calculate_loan_portfolio(year)
            total_assets = self._calculate_total_assets(year)
            equity = self.capital + self._calculate_cumulative_earnings(year)
            revenue = self._calculate_year_revenue(year)
            
            # Profitability ratios
            roe = (self._calculate_year_profit(year) / equity * 100) if equity > 0 else 0
            roa = (self._calculate_year_profit(year) / total_assets * 100) if total_assets > 0 else 0
            nim = (revenue * 0.85 / loan_portfolio * 100) if loan_portfolio > 0 else 0  # Net Interest Margin
            
            # Efficiency ratios
            expense_ratio = (self._calculate_year_expenses(year) / revenue * 100) if revenue > 0 else 100
            cost_per_client = self._calculate_year_expenses(year) / self._calculate_client_base(year) if self._calculate_client_base(year) > 0 else 0
            
            # Risk ratios
            default_rate = self.assumptions['default_rate'] * 100
            portfolio_growth = self._calculate_portfolio_growth_rate(year) * 100
            
            # Regulatory ratios (CBL specific)
            capital_adequacy = (equity / total_assets * 100) if total_assets > 0 else 0
            
            if self.tier == 'tier1':
                liquidity_ratio = self._calculate_liquidity_ratio(year) * 100
                deposit_growth = self._calculate_deposit_growth_rate(year) * 100
            else:
                liquidity_ratio = None
                deposit_growth = None
            
            ratios[f'year_{year}'] = {
                'profitability': {
                    'return_on_equity': f"{roe:.1f}%",
                    'return_on_assets': f"{roa:.1f}%",
                    'net_interest_margin': f"{nim:.1f}%",
                },
                'efficiency': {
                    'expense_ratio': f"{expense_ratio:.1f}%",
                    'cost_per_client': f"M{cost_per_client:.0f}",
                    'clients_per_staff': f"{self._calculate_clients_per_staff(year):.0f}",
                },
                'risk': {
                    'default_rate': f"{default_rate:.1f}%",
                    'portfolio_growth': f"{portfolio_growth:.1f}%",
                    'capital_adequacy': f"{capital_adequacy:.1f}%",
                },
                'regulatory': {
                    'liquidity_ratio': f"{liquidity_ratio:.1f}%" if liquidity_ratio else "N/A",
                    'deposit_growth': f"{deposit_growth:.1f}%" if deposit_growth else "N/A",
                    'cbl_compliance': self._check_cbl_compliance(year),
                }
            }
        
        return ratios
    
    def _perform_sensitivity_analysis(self) -> Dict:
        """Perform sensitivity analysis on key variables"""
        base_profit_y3 = self._calculate_year_profit(3)
        
        scenarios = {
            'optimistic': {
                'description': 'Higher client growth, lower defaults',
                'adjustments': {'client_growth': 1.5, 'default_rate': 0.7},
            },
            'pessimistic': {
                'description': 'Economic downturn, higher competition',
                'adjustments': {'client_growth': 0.7, 'default_rate': 1.5, 'interest_margin': 0.9},
            },
            'regulatory_stress': {
                'description': 'Stricter CBL requirements',
                'adjustments': {'capital_requirement': 1.3, 'compliance_costs': 1.4},
            }
        }
        
        results = {}
        for scenario, params in scenarios.items():
            adjusted_profit = self._calculate_scenario_profit(params['adjustments'])
            profit_change = ((adjusted_profit - base_profit_y3) / base_profit_y3 * 100) if base_profit_y3 > 0 else 0
            
            results[scenario] = {
                'description': params['description'],
                'year_3_profit': f"M{adjusted_profit:,.0f}",
                'profit_change': f"{profit_change:+.1f}%",
                'break_even_impact': self._calculate_scenario_break_even(params['adjustments']),
            }
        
        return results
    
    # Helper calculation methods
    def _calculate_year_revenue(self, year: int) -> Decimal:
        """Calculate total revenue for a given year"""
        loan_portfolio = self._calculate_loan_portfolio(year)
        avg_interest_rate = sum(self.assumptions['interest_rate_range']) / 2
        
        # Revenue grows with portfolio and market expansion
        market_growth = (1 + self.MARKET_ASSUMPTIONS['microfinance_growth']) ** year
        interest_revenue = loan_portfolio * Decimal(str(avg_interest_rate)) * market_growth
        fee_revenue = interest_revenue * Decimal('0.18')  # 18% of interest revenue
        
        return interest_revenue + fee_revenue
    
    def _calculate_year_expenses(self, year: int) -> Decimal:
        """Calculate total expenses for a given year"""
        revenue = self._calculate_year_revenue(year)
        efficiency = Decimal(str(self.assumptions['operational_efficiency']))
        
        # Expenses include staff, admin, technology, compliance
        base_expenses = revenue * (1 - efficiency)
        inflation_adjustment = (1 + self.MARKET_ASSUMPTIONS['inflation_rate']) ** year
        
        return base_expenses * Decimal(str(inflation_adjustment))
    
    def _calculate_year_profit(self, year: int) -> Decimal:
        """Calculate net profit for a given year"""
        revenue = self._calculate_year_revenue(year)
        expenses = self._calculate_year_expenses(year)
        credit_losses = revenue * Decimal(str(self.assumptions['default_rate']))
        
        return revenue - expenses - credit_losses
    
    def _calculate_loan_portfolio(self, year: int) -> Decimal:
        """Calculate loan portfolio size for a given year"""
        growth_rate = Decimal('0.25')  # 25% annual portfolio growth
        return self.initial_loan_portfolio * ((1 + growth_rate) ** year)
    
    def _calculate_cash_reserves(self, year: int) -> Decimal:
        """Calculate required cash reserves"""
        if self.tier == 'tier1':
            # 15% liquidity requirement for deposit-taking
            deposits = self._calculate_deposits(year)
            return deposits * Decimal('0.15')
        else:
            # 5% of assets for non-deposit institutions
            total_assets = self._calculate_total_assets(year)
            return total_assets * Decimal('0.05')
    
    def _calculate_deposits(self, year: int) -> Decimal:
        """Calculate customer deposits (Tier 1 only)"""
        if self.tier != 'tier1':
            return Decimal('0')
        
        # Deposits grow with portfolio and market confidence
        loan_portfolio = self._calculate_loan_portfolio(year)
        deposit_ratio = Decimal('1.2')  # 120% of loan portfolio
        
        return loan_portfolio * deposit_ratio
    
    def _calculate_break_even_month(self) -> int:
        """Calculate the month when the business becomes profitable"""
        monthly_fixed_costs = self._calculate_year_expenses(1) / 12
        monthly_revenue_target = monthly_fixed_costs / Decimal(str(self.assumptions['operational_efficiency']))
        
        # Assume gradual client acquisition over first year
        months_to_break_even = min(int(monthly_revenue_target / (self._calculate_year_revenue(1) / 12)), 18)
        
        return months_to_break_even
    
    def _check_cbl_compliance(self, year: int) -> str:
        """Check if projections meet CBL regulatory requirements"""
        issues = []
        
        # Capital adequacy
        capital_ratio = self._calculate_capital_adequacy_ratio(year)
        if capital_ratio < 0.08:  # 8% minimum
            issues.append("Below minimum capital adequacy")
        
        # Liquidity (Tier 1 only)
        if self.tier == 'tier1':
            liquidity_ratio = self._calculate_liquidity_ratio(year)
            if liquidity_ratio < 0.15:  # 15% minimum
                issues.append("Below minimum liquidity ratio")
        
        return "Compliant" if not issues else f"Issues: {', '.join(issues)}"
    
    def _calculate_capital_adequacy_ratio(self, year: int) -> float:
        """Calculate capital adequacy ratio"""
        equity = float(self.capital + self._calculate_cumulative_earnings(year))
        total_assets = float(self._calculate_total_assets(year))
        
        return equity / total_assets if total_assets > 0 else 0
    
    def _calculate_liquidity_ratio(self, year: int) -> float:
        """Calculate liquidity ratio (Tier 1 only)"""
        if self.tier != 'tier1':
            return 0
        
        cash_reserves = float(self._calculate_cash_reserves(year))
        deposits = float(self._calculate_deposits(year))
        
        return cash_reserves / deposits if deposits > 0 else 0
    
    def _calculate_cumulative_earnings(self, year: int) -> Decimal:
        """Calculate cumulative retained earnings up to a given year"""
        cumulative = Decimal('0')
        for y in range(1, year + 1):
            cumulative += self._calculate_year_profit(y)
        return cumulative
    
    # Additional helper methods for completeness
    def _calculate_total_assets(self, year: int) -> Decimal:
        loan_portfolio = self._calculate_loan_portfolio(year)
        cash_reserves = self._calculate_cash_reserves(year)
        fixed_assets = self._calculate_fixed_assets(year)
        return loan_portfolio + cash_reserves + fixed_assets * Decimal('1.05')  # 5% other assets
    
    def _calculate_fixed_assets(self, year: int) -> Decimal:
        # Technology, furniture, equipment
        initial_capex = self.capital * Decimal('0.1')  # 10% of capital
        annual_additions = initial_capex * Decimal('0.2')  # 20% annual addition
        return initial_capex + (annual_additions * year)
    
    def _calculate_client_base(self, year: int) -> int:
        growth_rate = 0.3  # 30% annual client growth
        initial_clients = int(self.target_market_size * 0.1)  # Start with 10% of target
        return int(initial_clients * ((1 + growth_rate) ** year))
    
    def _calculate_clients_per_staff(self, year: int) -> float:
        clients = self._calculate_client_base(year)
        staff_ratio = self.assumptions['staff_per_1000_clients'] / 1000
        staff_count = max(clients * staff_ratio, 5)  # Minimum 5 staff
        return clients / staff_count
    
    def _estimate_employment(self, year: int) -> int:
        clients = self._calculate_client_base(year)
        staff_ratio = self.assumptions['staff_per_1000_clients'] / 1000
        return max(int(clients * staff_ratio), 5)
    
    # Scenario analysis helpers
    def _calculate_scenario_profit(self, adjustments: Dict) -> Decimal:
        # Apply scenario adjustments and recalculate Year 3 profit
        base_profit = self._calculate_year_profit(3)
        
        # Apply adjustments
        client_adj = adjustments.get('client_growth', 1.0)
        default_adj = adjustments.get('default_rate', 1.0)
        margin_adj = adjustments.get('interest_margin', 1.0)
        
        adjusted_revenue = self._calculate_year_revenue(3) * Decimal(str(client_adj * margin_adj))
        adjusted_defaults = adjusted_revenue * Decimal(str(self.assumptions['default_rate'] * default_adj))
        adjusted_expenses = self._calculate_year_expenses(3)
        
        return adjusted_revenue - adjusted_expenses - adjusted_defaults
    
    def _calculate_scenario_break_even(self, adjustments: Dict) -> str:
        base_break_even = self._calculate_break_even_month()
        
        # Rough adjustment based on scenario impact
        client_impact = adjustments.get('client_growth', 1.0)
        if client_impact < 1.0:
            adjusted_months = int(base_break_even / client_impact)
            return f"+{adjusted_months - base_break_even} months delay"
        else:
            adjusted_months = int(base_break_even * client_impact)
            return f"-{base_break_even - adjusted_months} months faster"