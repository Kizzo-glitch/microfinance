"""
from docx import Document
from io import BytesIO
from django.utils import timezone

class CBLDocumentGenerator:

    def __init__(self, registration):
        self.registration = registration

    # ---------------------------------------------------------
    # AML Manual
    # ---------------------------------------------------------
    def aml_manual(self):
        doc = Document()
        doc.add_heading(f"{self.registration.company_name}", 0)
        doc.add_heading("AML/CFT Policy Manual", 1)
        doc.add_paragraph(f"Version 1.0 — {timezone.now().strftime('%B %Y')}")
        doc.add_page_break()
        doc.add_heading("1. INTRODUCTION", 1)
        doc.add_paragraph("This manual is prepared in compliance with CBL AML requirements...")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return ('aml_manual.docx', buffer)

    # ---------------------------------------------------------
    # Risk Manual
    # ---------------------------------------------------------
    def risk_manual(self):
        doc = Document()
        doc.add_heading(f"{self.registration.company_name}", 0)
        doc.add_heading("Risk Management Manual", 1)
        doc.add_paragraph("Prepared per CBL Risk Management Guidelines...")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return ('risk_manual.docx', buffer)

    # ---------------------------------------------------------
    # Complaints Procedure
    # ---------------------------------------------------------
    def complaints_procedure(self):
        doc = Document()
        doc.add_heading(f"{self.registration.company_name}", 0)
        doc.add_heading("Consumer Complaints and Redress Procedure", 1)
        doc.add_paragraph("This document covers procedures required under Regulation 11...")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return ('complaints_procedure.docx', buffer)

    # ---------------------------------------------------------
    # Business Plan
    # ---------------------------------------------------------
    def business_plan(self):
        doc = Document()
        doc.add_heading("BUSINESS PLAN", 0)
        doc.add_heading(self.registration.company_name, 1)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return ('business_plan.docx', buffer)

    # ---------------------------------------------------------
    # Cover Letter
    # ---------------------------------------------------------
    def cover_letter(self):
        doc = Document()
        doc.add_heading(self.registration.company_name, 0)
        doc.add_heading("Cover Letter", 1)
        doc.add_paragraph("To: The Governor, Central Bank of Lesotho...")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return ('cover_letter.docx', buffer)

    # ---------------------------------------------------------
    # Complete Pack
    # ---------------------------------------------------------
    def full_cbl_pack(self):
        return [
            self.cover_letter(),
            self.business_plan(),
            self.aml_manual(),
            self.risk_manual(),
            self.complaints_procedure(),
        ]
"""